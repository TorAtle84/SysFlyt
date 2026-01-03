from __future__ import annotations
import os
os.environ.setdefault("MPLBACKEND", "Agg")
try:
    import matplotlib
    matplotlib.use("Agg", force=True)
except Exception:
    pass

# --- Stdlib ---
import json
import re
import logging
import time
import shutil
import zipfile
import io
import threading
import secrets
import concurrent.futures
from pathlib import Path
from io import BytesIO
from collections import defaultdict, Counter

# --- 3rd party ---
import fitz
import extract_msg
import spacy
import joblib
import torch
import torch.nn.functional as F
import numpy as np
import sys
import subprocess
import pandas as pd
from pandas.errors import ParserError

from transformers import AutoTokenizer, AutoModel, AutoModelForSequenceClassification
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook
from openpyxl.styles import PatternFill, Alignment, Font, Border, Side
from openpyxl.utils import get_column_letter
from rapidfuzz.fuzz import ratio as fuzz_ratio, token_set_ratio, partial_ratio
from sentence_transformers import SentenceTransformer, util
from sklearn.metrics.pairwise import cosine_similarity
from celery import Celery

# --- App ---
from app.celery_instance import celery

# --- KS-2: Sorteringshjelpere ---
def _norm_key(s):
    try:
        return " ".join(str(s or "").split()).strip().lower()
    except Exception:
        return ""

def _score_val(x):
    # Prøver vanlige feltnavn for score/treffprosent
    for k in ("score", "treff", "treff_prosent", "Treff %", "match_score"):
        v = x.get(k)
        if v is not None:
            try:
                return float(v)
            except Exception:
                pass
    return -1.0

def _sort_requirements(reqs):
    # Primær: keyword alfabetisk (A–Å), sekundær: score synkende
    def _kw(x):
        for k in ("keyword", "søkeord", "sokeord", "key"):
            if k in x:
                return _norm_key(x.get(k))
        return _norm_key("")
    return sorted(reqs, key=lambda x: (_kw(x), -_score_val(x)))

def _sort_groups_order(keys):
    # Flytt 'Uspesifisert' (og ligatur-variant) til slutt
    unspec_idx = None
    cleaned = []
    for i, k in enumerate(keys):
        kl = _norm_key(k)
        if kl in ("uspesifisert", "uspesiﬁsert"):
            unspec_idx = i
        else:
            cleaned.append(k)
    if unspec_idx is not None:
        cleaned.append(keys[unspec_idx])
    return cleaned
# --- KS-2 slutt ---


# === Logging ===
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# === Stier ===
CURRENT_DIR = Path(__file__).resolve().parent
BASE_DIR = CURRENT_DIR.parent
TEMP_ROOT = CURRENT_DIR / "temp"
SYNONYM_PATH = CURRENT_DIR / "synonyms.json"
KEYWORD_PATH = CURRENT_DIR / "keywords.json"
MODEL_PATH = CURRENT_DIR / "data" / "krav_validator.pkl"
NS_CACHE_PATH = CURRENT_DIR / "data" / "ns_embeddings_cache.pkl"
FAG_PROFILER_PATH = CURRENT_DIR / "data" / "fag_profiler.pkl"
os.makedirs(TEMP_ROOT, exist_ok=True)

# === Kravsverb ===
GLOBAL_OBLIGATION_VERBS = {
    "skal","må","kreves","komplett","kreve","leveres","levere","installeres","monteres",
    "prosjekteres","sikre","utføres","dokumenteres","ber","inngår","består","utgjør",
    "påkrevd","obligatorisk","forutsatt","plikt","ansvarlig","forpliktet","sørge","bekrefter",
    "ivaretatt","fremgår","påse","dekke","godkjent","utarbeides","tilpasses","informerer",
    "fremmer","oversendes","anses","omfatter","bidrar","gjennomføre","avtales","benytte",
    "holde","stiller","tilbyr","sørger","pålagt","nødvendig","etterspurt","forventet",
    "påbudt","etterleve","implementere","overholde","sikret","bekreftet","iverksette",
    "tilrettelegge","gjennomføres","oppfylle","fullføre","måtte","opprettholde","etablere",
    "utskiftes","fornyes","besørge","legges","opprettholdes","klargjøres","demonteres",
    "fjernes","borttas","godkjenne","tillates","verifisere","kontrolleres","følges",
    "respektere","overleveres","overføres","tilbakeføres","reduseres","minimeres",
    "forhindres","unngås","opprette","registreres","protokollføres","arkiveres",
    "oppdateres","bevares","hindres","utredes","utbedres","etterse","ettergå",
    "aksepteres","autoriseres","vernes"
}

# === NB-BERT / MNLI init ===
DEVICE = torch.device("cuda" if torch.cuda.is_available() else "cpu")

NB_BERT_NAME = "NbAiLab/nb-bert-base"
MNLI_NAME = "NbAiLab/nb-bert-base-mnli"

nb_tokenizer = None
nb_model = None
mnli_tokenizer = None
mnli_model = None

try:
    nb_tokenizer = AutoTokenizer.from_pretrained(NB_BERT_NAME)
    nb_model = AutoModel.from_pretrained(NB_BERT_NAME).to(DEVICE).eval()
    NB_BERT_READY = True
    logging.info("NB-BERT lastet.")
except Exception as e:
    NB_BERT_READY = False
    logging.warning(f"Kunne ikke laste NB-BERT ({NB_BERT_NAME}): {e}")

try:
    mnli_tokenizer = AutoTokenizer.from_pretrained(MNLI_NAME)
    mnli_model = AutoModelForSequenceClassification.from_pretrained(MNLI_NAME).to(DEVICE).eval()
    MNLI_READY = True
    logging.info("NB-MNLI lastet.")
except Exception as e:
    MNLI_READY = False
    logging.warning(f"Kunne ikke laste NB-MNLI ({MNLI_NAME}): {e}")

def _mean_pooling(model_output, attention_mask):
    token_embeddings = model_output[0]
    input_mask_expanded = attention_mask.unsqueeze(-1).expand(token_embeddings.size()).float()
    return (token_embeddings * input_mask_expanded).sum(1) / input_mask_expanded.sum(1).clamp(min=1e-9)

@torch.no_grad()
def nb_bert_encode(texts: list[str]) -> torch.Tensor:
    """
    Returnerer L2-normaliserte setningsembeddinger for norske tekster.
    Kaster RuntimeError hvis modellen ikke er klar.
    """
    if not NB_BERT_READY or nb_tokenizer is None or nb_model is None:
        raise RuntimeError("NB-BERT ikke initialisert.")
    tokens = nb_tokenizer(texts, padding=True, truncation=True, return_tensors="pt")
    tokens = {k: v.to(DEVICE) for k, v in tokens.items()}
    outputs = nb_model(**tokens)
    sent_emb = _mean_pooling(outputs, tokens['attention_mask'])
    sent_emb = F.normalize(sent_emb, p=2, dim=1)
    return sent_emb

@torch.no_grad()
def nb_mnli_predict(premise: str, hypothesis: str) -> dict:
    """
    Returnerer sannsynligheter for {'contradiction','neutral','entailment'}.
    Kaster RuntimeError hvis modellen ikke er klar.
    """
    if not MNLI_READY or mnli_tokenizer is None or mnli_model is None:
        raise RuntimeError("NB-MNLI ikke initialisert.")
    tokens = mnli_tokenizer(premise, hypothesis, return_tensors="pt", truncation=True, padding=True)
    tokens = {k: v.to(DEVICE) for k, v in tokens.items()}
    logits = mnli_model(**tokens).logits
    probs = torch.softmax(logits, dim=-1)[0].tolist()
    # Label-rekkefølge: [contradiction, neutral, entailment]
    return {"contradiction": probs[0], "neutral": probs[1], "entailment": probs[2]}   
    
# --- Domeneprofiler (per fag) ---
DOMAIN_PROFILES = {
    # Fallback når vi ikke vet fag: bred, men trygg
    "generic": {
        "aliases": [],
        "terms": set(),
        # tall + generelle enheter/tegn, inkl. minusgrader og dB(A)
        "units_re": re.compile(
            r"(\b-?\d+[.,]?\d*\s*(?:ppm|%|°\s*c|k?w|w|k?va|var|db(?:\s*\(a\))?|pa|kpa)\b)"
            r"|(\b[<>]=?\s*-?\d+[.,]?\d*\b)",
            re.IGNORECASE
        ),
    },

    # ——— VENTILASJON ———
    "ventilasjon": {
        "aliases": [
            "ventilasjon","ventilasjonsanlegg","tilluft","avtrekk","overstrøm",
            "vav","vav-spjeld","cav","dcv",
            "undertrykk","overtrykk","trykksetting","trykksatt","trykksettes",
            "dut","dimensjonerende utetemperatur","utetemperatur",
            "tilluftstemperatur","operativ temperatur","romtemperatur"
        ],
        "terms": {
            "ventilasjon","ventilasjonsanlegg","tilluft","avtrekk","overstrøm",
            "vav","vav-spjeld","cav","dcv",
            "sfp","kW/(m³/s)","k W/(m ³ / s)","w/(m³/s)","m³/s","m3/s","m ³ / s",
            "m³/m²h","m3/m2h","m³/m²/t","m3/m2/t","l/s per person","l/s pr person",
            "co2","co₂","ppm","pa","kpa","trykkfall","trykkdifferanse","differansetrykk",
            "undertrykk","overtrykk","trykksetting","trykksatt","trykksettes",
            "dut","dimensjonerende utetemperatur","utetemperatur",
            "tilluftstemperatur","operativ temperatur","romtemperatur","°c","grader",
            "varmegjenvinning","temperaturvirkningsgrad","bypass-spjeld","lydfelle",
            "innregulering","ns-en 12599","ns-en 16798","røykspjeld","avkast","inntak",
            "sluse","sluser","toalettsoner","operasjonsrom","isoleringsrom","trapperom",
            "filterklasse epm1","en iso 16890"
        },
        "units_re": re.compile(
            r"(\b-?\d+[.,]?\d*\s*(?:"
            r"m\s*[³3]|m\s*[²2]|"                       # m³ / m²
            r"m\s*[³3]\s*/\s*s|"                        # m³/s
            r"m\s*[³3]\s*/\s*m\s*[²2]\s*/\s*t|"         # m³/m²/t
            r"l/s|l/min|m³/h|"                           # luftmengder
            r"ppm|k\s*w|kw|w|°\s*c|pa|kpa"               # CO₂/energi/temp/trykk
            r")\b)"
            r"|(\b[<>]=?\s*-?\d+[.,]?\d*\b)",
            re.IGNORECASE
        ),
    },

    # ——— ELEKTRO ———
    "elektro": {
        "aliases": [
            "elektro","el","strøm","belysning","lys","brannalarm","nødlys","adgang",
            "ik","ip","ups","tavle","kurs","fordeling","nek 400"
        ],
        "terms": {
            "elektro","strøm","kabel","kurs","gruppe","vern","rcd","rcbo","spn","spd","selektivitet",
            "belysning","lux","lysnivå","armatur","dali","dag/natt","tilstedeværelse",
            "brannalarm","detektor","varslingsanlegg","nødlys","ledelys",
            "adgangskontroll","kortleser","porttelefon","fiber","nettverk","ups",
            "hovedtavle","underfordeling","tn-s","it nett","ip","ik","nek 400","feilstrøm",
            "hz","v","kv","a","ka","ω","ohm","mm²","cos φ","ik","ik3","z loop"
        },
        "units_re": re.compile(
            r"(\b\d+[.,]?\d*\s*(?:v|kv|a|ka|w|kw|kva|hz|lux|db(?:\s*\(a\))?|mm²|mm2|ω|ohm)\b)"
            r"|(\bip\d{2}\b)|(\bik\d{2}\b)|(\bik\s*\d+[.,]?\d*\s*k?a\b)",
            re.IGNORECASE
        ),
    },

    # ——— RØRLEGGER (Sanitær/Sprinkler) ———
    "rørlegger": {
        "aliases": [
            "rør","vvs","sanitær","sprinkler","vann","avløp","bereder","pumpe","trykktest","legionella"
        ],
        "terms": {
            "vann","avløp","spillvann","gråvann","sanitær","rør","pe-x","pex","cu","rustfritt",
            "sirkulasjon","pumpe","sirkulasjonspumpe","fettutskiller","bereder","varmtvann","blandearmatur",
            "sprinkler","sprinklersentral","alarmventil","sprinklerhode","ns-en 12845",
            "lekkasje","lekkasjesikring","trykktest","tetthetsprøve","legionella","tilbakestrømssikring",
            "forbruksvann","trykklasse","kjel","varmesentral","vannmåler","reduksjonsventil","filter",
            "dn","dimensjon dn","rør-i-rør"
        },
        "units_re": re.compile(
            r"(\b-?\d+[.,]?\d*\s*(?:l/s|l/min|m³/h|bar|kpa|pa|°\s*c)\b)"
            r"|(\bdn\s*\d+\b)"
            r"|(\b[<>]=?\s*-?\d+[.,]?\d*\b)",
            re.IGNORECASE
        ),
    },

    # ——— BYGGAUTOMASJON (BAS/BMS/SD) ———
    "byggautomasjon": {
        "aliases": [
            "sd","bms","bas","byggautomasjon","bacnet","modbus","knx","lon","trend","romkontroller"
        ],
        "terms": {
            "byggautomasjon","sd","bms","bacnet","modbus","knx","lon","opc ua","mqtt",
            "io","i/o","di","do","ai","ao","pid","regulator","sekvensstyring",
            "setpunkt","romkontroller","trend","trendlogg","historikk","rapport","alarmer","hendelser",
            "energioppfølging","forbrukslogg","integrasjon","gateway","api","rollebasing"
        },
        "units_re": re.compile(
            r"(\b-?\d+[.,]?\d*\s*(?:°\s*c|%|ppm|kwh|wh|s|min|ms)\b)"
            r"|(\b[<>]=?\s*-?\d+[.,]?\d*\b)",
            re.IGNORECASE
        ),
    },

    # ——— KULDE / KJØLING / VARMEPUMPER ———
    "kulde": {
        "aliases": [
            "kulde","kjøleanlegg","chiller","dx","varmepumpe","f-gass","kuldemedium","brine","glykol"
        ],
        "terms": {
            "f-gass","kuldemedium","r32","r410a","r1234ze","r290","r744","co2 kjøling",
            "kompressor","kondensator","fordamper","ekspansjonsventil","hetgass","underkjøling",
            "kondensvann","brine","glykol","plateveksler",
            "cop","eer","seer","scop","kw kjølekapasitet","kw varmeeffekt","Δt","°c",
            "lekkasjekontroll","tetthetskontroll","gwp"
        },
        "units_re": re.compile(
            r"(\b-?\d+[.,]?\d*\s*(?:kw|w|°\s*c|kpa|pa)\b)"
            r"|(\b(cop|eer|seer|scop)\s*[>=]?\s*\d+[.,]?\d*\b)"
            r"|(\b[<>]=?\s*-?\d+[.,]?\d*\b)",
            re.IGNORECASE
        ),
    },

    # ——— TOTALENTREPRENØR / KONTRAKT ———
    "totalentreprenør": {
        "aliases": ["totalentreprenør","entreprenør","hovedentreprenør","te"],
        "terms": {
            "ns 8407","framdriftsplan","sha-plan","kvalitetssystem","kontrollplan",
            "as built","fdv","prøvedrift","overtakelse","endringsordre","tilleggsarbeid",
            "rigg og drift","avviksbehandling","kompetansekrav","bemanningsplan","dokumentasjon"
        },
        "units_re": re.compile(
            r"(\b\d+[.,]?\d*\s*(?:%|dager|uke|uker|måneder|timer|kr)\b)"
            r"|(\b[<>]=?\s*\d+[.,]?\d*\b)",
            re.IGNORECASE
        ),
    },

    # ——— BYGGHERRE / BESTILLER ———
    "byggherre": {
        "aliases": ["byggherre","bh","oppdragsgiver","prosjekteier"],
        "terms": {
            "kravdokument","funksjonskrav","romprogram","arealprogram","brukerutstyr",
            "beslutningspunkt","tilvalg","opsjon","møtereferat","endringshåndtering",
            "overtakelsesprotokoll","godkjenne","beslutte"
        },
        "units_re": re.compile(
            r"(\b\d+[.,]?\d*\s*(?:%|dager|uke|uker|måneder|kr)\b)"
            r"|(\b[<>]=?\s*\d+[.,]?\d*\b)",
            re.IGNORECASE
        ),
    },

    # ——— ØKONOMI ———
    "økonomi": {
        "aliases": ["økonomi","kontrakt","kostnad","betaling","pris","fakturaplan"],
        "terms": {
            "dagmulkt","vederlag","opsjon","opsjonspris","indeksregulering","ns 3420",
            "fakturaplan","enhetspris","timepris","endringsordre","kontraktssum",
            "prisoverslag","tilbudssum","milepæl","betalingsplan"
        },
        "units_re": re.compile(
            r"(\b\d+[.,]?\d*\s*(?:%|kr|nok|eur|usd)\b)"
            r"|(\b[<>]=?\s*\d+[.,]?\d*\b)",
            re.IGNORECASE
        ),
    },
}


UNITS_REGEX = re.compile(
    r"(\b-?\d+[.,]?\d*\s*(?:"
    r"m\s*[³3]|m\s*[²2]|"                      # m³ / m²
    r"m\s*[³3]\s*/\s*s|"                       # m³/s
    r"m\s*[³3]\s*/\s*m\s*[²2]\s*/\s*t|"        # m³/m²/t
    r"l/s|l/min|m³/h|"                          # vanlige VVS-enheter
    r"ppm|k\s*w|kw|w|°\s*c|pa|kpa"              # °C (med/uten mellomrom), Pa/kPa, kW/W
    r")\b)|(\b[<>]=?\s*-?\d+[.,]?\d*\b)",
    flags=re.IGNORECASE
)

# === Tekstrens ===
def clean_text(text: str) -> str:
    cleaned_text = re.sub(r'[\x00-\x1F\x7F]', '', text)
    cleaned_text = cleaned_text.replace('\r\n', '\n').replace('\r', '\n')
    cleaned_text = re.sub(r'(Fra:|Sendt:|Til:|Kopi:|Emne:)\s*', r'\n\1 ', cleaned_text, flags=re.IGNORECASE)
    cleaned_text = re.sub(r'(\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b|\bhttps?://\S+|\bwww\.\S+)', r'\1\n', cleaned_text, flags=re.IGNORECASE)
    cleaned_text = re.sub(r'(P\.O\.Box|Telephone|Mobile|Tlf:)\s*', r'\n\1 ', cleaned_text, flags=re.IGNORECASE)
    cleaned_text = re.sub(r'^\s*([._-]{3,}|\*{3,}|={3,}|#{3,})\s*$', '', cleaned_text, flags=re.MULTILINE)
    cleaned_text = re.sub(r'^\s*(\d+\s*){1,5}$', '', cleaned_text, flags=re.MULTILINE)
    cleaned_text = re.sub(r'(\S)\s*(\.{2,}|\-{2,})\s*(\d+)', r'\1 \3', cleaned_text)
    cleaned_text = re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', cleaned_text, flags=re.IGNORECASE)
    cleaned_text = re.sub(r'(?<!\n)\n(?!\n)', ' ', cleaned_text)
    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
    cleaned_text = re.sub(r'([a-zæøå])([A-ZÆØÅ])', r'\1 \2', cleaned_text)
    cleaned_text = re.sub(r'\s+([,.?!;])', r'\1', cleaned_text)
    cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
    return cleaned_text

# === Korttekst ===
def _generate_short_text(text: str, max_len: int = 120) -> str:
    """Kort, forklarende oppsummering av kravet (DUT, SFP, Pa, CO₂, °C, m³/s, osv.)."""
    if not text:
        return ""
    t = re.sub(r"\s+", " ", text).strip()
    parts = []

    # DUT / dimensjonerende utetemperatur
    m = re.search(r"(DUT|dimensjonerende|dimensjonerende)\s+utetemperatur[^.;]*?(-?\d+[.,]?\d*)\s*°\s*C", t, re.I)
    if m:
        parts.append(f"DUT {m.group(2)} °C")

    # Minimum utetemperatur
    m = re.search(r"min(?:imum)?\s*(?:ute-?\s*temperatur|utetemperatur)[^.;]*?(-?\d+[.,]?\d*)\s*°\s*C", t, re.I)
    if m:
        parts.append(f"Min. utetemp {m.group(1)} °C")

    # Maks tilluftstemperatur
    m = re.search(r"maks(?:imal)?\s*(?:tillufts-?\s*temperatur|tilluftstemperatur)[^.;]*?(\d+[.,]?\d*)\s*°\s*C", t, re.I)
    if m:
        parts.append(f"Maks tilluft {m.group(1)} °C")

    # Operativ/romtemperatur
    m = re.search(r"(operativ temperatur|romtemperatur)[^.;]*?(-?\d+[.,]?\d*)\s*°\s*C(?:[^.;]*?(\d+[.,]?\d*)\s*°\s*C)?", t, re.I)
    if m:
        if m.group(3):
            parts.append(f"{m.group(1).title()} {m.group(2)}–{m.group(3)} °C")
        else:
            parts.append(f"{m.group(1).title()} {m.group(2)} °C")

    # Trykkdifferanse / under-/overtrykk
    m = re.search(r"((?:under|over)trykk|trykk(?:differanse|setting))[^.;]*?(-?\d+[.,]?\d*)\s*(k?\s*Pa)", t, re.I)
    if m:
        parts.append(re.sub(r"\s+", " ", m.group(0)).strip(" .;"))

    # CO₂
    m = re.search(r"CO[2₂][^.;]*?(\d+[.,]?\d*)\s*ppm", t, re.I)
    if m:
        parts.append(f"CO₂ {m.group(1)} ppm")

    # SFP
    m = re.search(r"\bSFP\b[^.;]*?([<>]=?)\s*(\d+[.,]?\d*)\s*(k?\s*W)\s*/\s*\(?\s*m\s*[³3]\s*/\s*s\)?", t, re.I)
    if m:
        parts.append(f"SFP {m.group(1)} {m.group(2)} kW/(m³/s)")

    # Varmegjenvinning
    m = re.search(r"varmegjenvinning[^.;]*?([<>]=?)\s*(\d{1,3})\s*%", t, re.I)
    if m:
        parts.append(f"Varmegjenvinning {m.group(1)} {m.group(2)} %")

    # Luftmengde
    m = re.search(r"(\d+[.,]?\d*)\s*m\s*[³3]\s*/\s*m\s*[²2]\s*/\s*t", t, re.I)
    if m:
        m2 = re.search(r"[^.;]*\d+[.,]?\d*\s*m\s*[³3]\s*/\s*m\s*[²2]\s*/\s*t[^.;]*", t, re.I)
        if m2:
            parts.append(re.sub(r"\s+", " ", m2.group(0)).strip(" .;"))

    # VAV / sekvensregulering
    m = re.search(r"\bVAV[^.;]*", t, re.I)
    if m:
        parts.append(re.sub(r"\s+", " ", m.group(0)).strip(" .;"))
    m = re.search(r"\bsekvensregulering[^.;]*", t, re.I)
    if m:
        parts.append(re.sub(r"\s+", " ", m.group(0)).strip(" .;"))

    parts = [p for p in parts if p]
    parts = list(dict.fromkeys(parts))
    summary = "; ".join(parts) if parts else t.split(".")[0]
    summary = summary.strip()
    return summary[:max_len] + ("..." if len(summary) > max_len else "")

# === Foreslå nøkkelord (C1) ===
def _extract_suggested_keywords(all_krav_texts: list, existing_keywords: list, existing_synonyms: dict, top_n: int = 10) -> list:
    if not nlp or not all_krav_texts:
        return []
    combined_text = " ".join(all_krav_texts)
    doc_combined = nlp(combined_text)
    existing_terms_lower = set(kw.lower() for kw in existing_keywords)
    for base_kw, syns_list in (existing_synonyms or {}).items():
        existing_terms_lower.add(base_kw.lower())
        existing_terms_lower.update(s.lower() for s in syns_list)
    suggested_terms_counter = Counter()
    for chunk in doc_combined.noun_chunks:
        chunk_text = chunk.text.strip().lower()
        if (chunk_text and len(chunk_text.split()) <= 4 and
            chunk_text not in existing_terms_lower and
            not any(token.is_stop for token in chunk) and
            not re.fullmatch(r'[\s\W]+', chunk_text)):
            suggested_terms_counter[chunk_text] += 1
    for token in doc_combined:
        lemma = token.lemma_.lower()
        if not token.is_stop and token.is_alpha and lemma not in existing_terms_lower and token.pos_ in ["NOUN","ADJ","VERB"]:
            suggested_terms_counter[lemma] += 1
    final = [term for term, _ in suggested_terms_counter.most_common(top_n*2)
             if len(term) > 2 and term.count(' ') < 3 and term not in existing_terms_lower]
    return list(dict.fromkeys(final))[:top_n]

# === NS-cache (B1) ===
def _load_ns_cache(cache_path: Path, pdf_standards_config: dict) -> dict | None:
    if not cache_path.exists():
        return None
    cache_mtime = os.path.getmtime(cache_path)
    for _, details in pdf_standards_config.items():
        if details.get("aktiv", False):
            ns_pdf_path = Path(details["path"])
            if not ns_pdf_path.exists() or os.path.getmtime(ns_pdf_path) > cache_mtime:
                return None
    try:
        with open(cache_path, 'rb') as f:
            return joblib.load(f)
    except Exception:
        return None

def _save_ns_cache(data: dict, cache_path: Path):
    try:
        with open(cache_path, 'wb') as f:
            joblib.dump(data, f)
    except Exception as e:
        logging.error(f"Kunne ikke lagre NS-cache: {e}")

# === Last modeller ===
logging.info("Laster NLP/ML-modeller...")
nlp = None
try:
    nlp = spacy.load("nb_core_news_lg")
    logging.info("SpaCy 'nb_core_news_lg' lastet.")
except Exception as e:
    logging.error(f"SpaCy feilet: {e}")
    nlp = None

krav_validator_model = None
if MODEL_PATH.exists():
    try:
        krav_validator_model = joblib.load(MODEL_PATH)
        logging.info(f"ML-modell lastet: {MODEL_PATH}")
    except Exception as e:
        logging.error(f"ML-modell feilet: {e}")
else:
    logging.error(f"ML-modell ikke funnet: {MODEL_PATH}")

semantic_model = None
try:
    semantic_model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
    logging.info("SentenceTransformer lastet.")
except Exception as e:
    logging.error(f"Semantisk modell feilet: {e}")
    semantic_model = None

fag_model = None
if FAG_PROFILER_PATH.exists():
    try:
        fag_model = joblib.load(FAG_PROFILER_PATH)
        logging.info(f"Fag-modell lastet: {FAG_PROFILER_PATH}")
    except Exception as e:
        logging.error(f"Fag-modell feilet: {e}")
else:
    logging.error(f"Fag-modell ikke funnet: {FAG_PROFILER_PATH}")

# === Klassifisering / farger ===
regex_patterns = {
    "installasjon":[r"\b(montere|installer|kobles til|settes opp|montasje|tetting|hull|påstøpes|festes|ankres|forbindes|oppheng)\b"],
    "leveranse":[r"\b(leveres|inkludert|medfølger|omfatter|bestilles|komplett|tilbys|tilgjengelig)\b"],
    "dokumentasjon":[r"\b(fdv|manual|datablad|protokoll|dokumentasjon|breeam|ns3935|ns6450|brukerveiledning|sertifikat|opplæring)\b"],
    "miljø/energi":[r"\b(breeam|energi(?:måler|beregning)|epc|miljø(?:profil|rapport)|co2|klimagass|energieffektivisering)\b"],
    "opsjon":[r"\b(valgfritt|opsjon|kan leveres|ekstra|opsjonelt|mulighet for)\b"]
}
KRAVTYPE_COLORS = {"installasjon":"FFFF99","leveranse":"CCFFCC","dokumentasjon":"CCCCFF","opsjon":"FFCCFF","krav":"DDDDDD","generell":"DED8EB","miljø/energi":"FFC285"}
GROUP_COLORS = {"byggherre":"FFD966","totalentreprenør":"FCE5CD","elektro":"FFF2CC","ventilasjon":"D0E0E3","rørlegger":"EAD1DC","byggautomasjon":"CCE5FF","prosjektering":"D9D2E9","økonomi":"FFF0F5","elektro, prosjektering":"FCE5CD","ventilasjon, prosjektering":"D9D2E9","rørlegger, prosjektering":"F4CCCC","byggautomasjon, prosjektering":"D9EAD3","elektro, økonomi":"FAFAD2","ventilasjon, økonomi":"E0FFFF","rørlegger, økonomi":"FFE4E1","byggautomasjon, økonomi":"F0FFF0","Uspesifisert":"EEEEEE","GK":"D9EAD3","Bravida":"CFE2F3","Caverion":"FCE5CD","Vestrheim":"F4CCCC"}
PDF_STANDARDER = {
    "NS8415":{"path": str(CURRENT_DIR / "data/NS8415.pdf"), "aktiv": True},
    "NS8417":{"path": str(CURRENT_DIR / "data/NS8417.pdf"), "aktiv": True}
}

# === Indekser NS-standarder (med cache) ===
global_ns_data = defaultdict(list)
cached_ns = None
if nlp and semantic_model:
    cached_ns = _load_ns_cache(NS_CACHE_PATH, PDF_STANDARDER)
if cached_ns is not None:
    global_ns_data.update(cached_ns)
else:
    if nlp and semantic_model:
        for std, details in PDF_STANDARDER.items():
            if not details.get("aktiv", False):
                continue
            p = Path(details["path"])
            if not p.exists():
                logging.warning(f"NS-fil mangler: {p}")
                continue
            try:
                with fitz.open(p) as doc:
                    for page_num, page in enumerate(doc):
                        page_text = clean_text(page.get_text("text") or "")
                        if not page_text:
                            continue
                        ns_doc = nlp(page_text)
                        for s in ns_doc.sents:
                            stext = s.text.strip()
                            if len(stext.split()) > 3 and not re.fullmatch(r'[\s\W]*', stext):
                                emb = semantic_model.encode(stext, convert_to_tensor=True)
                                global_ns_data[std].append({"side": page_num+1, "tekst": stext, "embedding": emb})
                logging.info(f"Indeksert {std}: {len(global_ns_data[std])} segmenter")
            except Exception as e:
                logging.error(f"NS-indeksering feilet for {std}: {e}", exc_info=True)
        if global_ns_data:
            _save_ns_cache(global_ns_data, NS_CACHE_PATH)

# === Diverse hjelpere ===
def classify_group_ai(text: str):
    if not fag_model or not text.strip():
        return "Uspesifisert", []
    try:
        # Format A: forventa av dagens kode
        if all(k in fag_model for k in ("vectorizer","matrix","classes")):
            vec = fag_model["vectorizer"].transform([text])
            if vec.shape[1] != fag_model["matrix"].shape[1]:
                logging.error("[classify_group_ai] Dimensjonsmismatch vectorizer/matrix.")
                return "Uspesifisert", []
            sims = cosine_similarity(vec, fag_model["matrix"]).ravel()
            ranked = sorted(zip(fag_model["classes"], sims), key=lambda x: x[1], reverse=True)
            best = ranked[0][0] if ranked and ranked[0][1] > 0.1 else "Uspesifisert"
            return best, ranked

        # Format B: artefakt fra train_fag_classifier.py (pipeline + labels)
        if all(k in fag_model for k in ("pipeline","labels")):
            pipe = fag_model["pipeline"]
            labels = list(fag_model.get("labels", []))
            # bruk beslutningsfunksjon/sannsynlighet som “score”
            try:
                scores = getattr(pipe, "decision_function", pipe.predict_proba)([text])[0]
            except Exception:
                scores = pipe.predict_proba([text])[0]
            # sørg for numpy-array
            scores = np.array(scores).ravel()
            ranked = sorted(zip(labels, scores), key=lambda x: x[1], reverse=True)
            best = ranked[0][0] if ranked else "Uspesifisert"
            return best, ranked

        logging.error("[classify_group_ai] Uventet fag_model-format.")
        return "Uspesifisert", []
    except Exception:
        logging.error("[classify_group_ai] Feil under klassifisering.", exc_info=True)
        return "Uspesifisert", []

def is_valid_requirement(tekst: str) -> bool:
    if not krav_validator_model:
        return len(tekst.split()) >= 4
    if len(tekst.split()) < 4:
        return False
    try:
        pred = krav_validator_model.predict([tekst])
        return pred[0] == 1
    except Exception:
        return True

def classify_type(tekst):
    tl = tekst.lower()
    for kravtype, patterns in regex_patterns.items():
        for pattern in patterns:
            if re.search(pattern, tl):
                return kravtype
    return "krav"

def classify_group(tekst):
    tl = tekst.lower()
    if "byggherre" in tl: return "byggherre"
    if "totalentreprenør" in tl: return "totalentreprenør"
    if any(w in tl for w in ["økonomi","kostnad","pris","dagmulkt"]):
        if any(w in tl for w in ["elektro","strøm","kabel"]): return "elektro, økonomi"
        if any(w in tl for w in ["ventilasjon","luft"]): return "ventilasjon, økonomi"
        if any(w in tl for w in ["rør","sanitær","vann"]): return "rørlegger, økonomi"
        if any(w in tl for w in ["sd","bas","byggautomasjon"]): return "byggautomasjon, økonomi"
        return "økonomi"
    if "prosjektering" in tl or "detaljprosjektering" in tl:
        if any(w in tl for w in ["elektro","strøm","kabel"]): return "elektro, prosjektering"
        if any(w in tl for w in ["ventilasjon","luft"]): return "ventilasjon, prosjektering"
        if any(w in tl for w in ["rør","sanitær","vann"]): return "rørlegger, prosjektering"
        if any(w in tl for w in ["sd","bas","byggautomasjon"]): return "byggautomasjon, prosjektering"
        return "prosjektering"
    if any(w in tl for w in ["strøm","kabel","elektro","sikring","belysning","brannalarm","adgangskontroll","lys","ups","tavle","solcelle","energimåler","nettverk","mast","port","telefon"]): return "elektro"
    if any(w in tl for w in ["luft","ventilasjon","avtrekk","tilluft","vav","cav","overstrøm","ventilering","brannspjeld","hette","rist","kanal","batteriaggregat","røykavtrekk","blikk","inneklima","sfp","varmegjenvinning","kjølegjenvinning"]): return "ventilasjon"
    if any(w in tl for w in ["vann","avløp","sanitær","rør","sirkulasjon","pumpe","fettutskiller","bereder","borrehull","lekkasje","sprinkler","kaldras"]): return "rørlegger"
    if any(w in tl for w in ["sd","bas","byggautomasjon","kontroller","sd-anlegg","program","regulering","bus","bacnet","modbus","autonom"]): return "byggautomasjon"
    if "gk" in tl: return "GK"
    if "bravida" in tl: return "Bravida"
    if "caverion" in tl: return "Caverion"
    if "vestrheim" in tl: return "Vestrheim"
    return "Uspesifisert"

def delayed_cleanup(path, delay_seconds=300):
    def _cleanup():
        time.sleep(delay_seconds)
        try:
            if os.path.exists(path):
                shutil.rmtree(path)
                logging.info(f"[Cleanup] Slettet temp-mappe etter {delay_seconds}s: {path}")
        except Exception as e:
            logging.warning(f"[Cleanup] Klarte ikke slette {path}: {e}", exc_info=True)
    threading.Thread(target=_cleanup, daemon=True).start()

def _convert_doc_to_docx(input_path: Path, out_dir: Path) -> Path | None:
    """
    Konverterer .doc → .docx ved hjelp av LibreOffice (soffice).
    Returnerer sti til .docx ved suksess, ellers None.
    """
    try:
        out_dir.mkdir(parents=True, exist_ok=True)
        # LibreOffice skriver til out_dir
        cmd = [
            "soffice", "--headless", "--convert-to", "docx",
            "--outdir", str(out_dir), str(input_path)
        ]
        res = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if res.returncode != 0:
            logging.warning(f"[.doc→.docx] Konvertering feilet for {input_path.name}: {res.stderr or res.stdout}")
            return None
        out_path = out_dir / (input_path.stem + ".docx")
        return out_path if out_path.exists() else None
    except FileNotFoundError:
        logging.warning("[.doc→.docx] Fant ikke 'soffice' (LibreOffice). Installer LibreOffice i miljøet for .doc-støtte.")
        return None
    except Exception as e:
        logging.warning(f"[.doc→.docx] Uforutsett feil for {input_path.name}: {e}", exc_info=True)
        return None

def _extract_text_from_docx_bytes(data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        logging.warning(f"[DOCX-bytes] Kunne ikke lese DOCX-bytes: {e}")
        return ""

# === Utheving (dynamisk synonymlasting) ===
def _load_word_family_map():
    m = {}
    if SYNONYM_PATH.exists():
        try:
            with open(SYNONYM_PATH, "r", encoding="utf-8") as f:
                syns = json.load(f)
            for base, syn_list in syns.items():
                fam = [base.lower()] + [s.lower() for s in syn_list]
                for w in fam:
                    m[w] = fam
        except json.JSONDecodeError:
            pass
    return m

def highlight_docx(path, keywords, output):
    word_map = _load_word_family_map()
    doc = Document(path)
    keywords_lower = [k.lower() for k in keywords]
    all_terms = set(keywords_lower)
    for base in keywords_lower:
        if base in word_map: all_terms.update(word_map[base])
    terms = sorted([re.escape(t) for t in all_terms if t], key=len, reverse=True)
    if not terms:
        doc.save(output); return
    pattern = r'\b(' + '|'.join(terms) + r')\b'
    for p in doc.paragraphs:
        if not p.text.strip(): continue
        rebuilt, last = [], 0
        for m in re.finditer(pattern, p.text, flags=re.IGNORECASE):
            if m.start() > last: rebuilt.append((p.text[last:m.start()], False))
            rebuilt.append((m.group(0), True)); last = m.end()
        if last < len(p.text): rebuilt.append((p.text[last:], False))
        while p.runs:
            r = p.runs[0]
            r._r.getparent().remove(r._r)

        # Bygg opp nye runs
        for txt, hi in rebuilt:
            if not txt:
                continue
            run = p.add_run(txt)
            if hi:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.save(output)

def highlight_pdf(path, keywords, output):
    word_map = _load_word_family_map()
    doc = fitz.open(path)
    keywords_lower = [k.lower() for k in keywords]
    all_terms = set(keywords_lower)
    for base in keywords_lower:
        if base in word_map: all_terms.update(word_map[base])
    for page in doc:
        for kw in all_terms:
            for inst in page.search_for(kw, flags=1):
                page.add_highlight_annot(inst).update()
    doc.save(output, garbage=4, deflate=True, clean=True)
    doc.close()

def highlight_excel(path, keywords, output):
    word_map = _load_word_family_map()
    wb = load_workbook(path)
    fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
    keywords_lower = [k.lower() for k in keywords]
    all_terms = set(keywords_lower)
    for base in keywords_lower:
        if base in word_map: all_terms.update(word_map[base])
    for s in wb.worksheets:
        for row in s.iter_rows():
            for c in row:
                if c.value is not None:
                    if any(term in str(c.value).lower() for term in all_terms):
                        c.fill = fill
    wb.save(output)

def highlight_text(path, keywords, output):
    word_map = _load_word_family_map()
    lines = Path(path).read_text(encoding='utf-8').splitlines()
    keywords_lower = [k.lower() for k in keywords]
    all_terms = set(keywords_lower)
    for base in keywords_lower:
        if base in word_map: all_terms.update(word_map[base])
    with open(output, "w", encoding="utf-8") as f:
        for line in lines:
            if any(term in line.lower() for term in all_terms):
                f.write(f"* {line.strip()}\n")
            else:
                f.write(f"{line.strip()}\n")

# === Excel-rapport (FLAT / legacy) ===
def generate_total_excel_report(
    all_requirements,
    keywords,
    total_files_scanned: int,
    total_requirements_found: int,
    keyword_counts: defaultdict,
    synonym_counts: defaultdict,
    semantic_search_counts: defaultdict
):
    wb = Workbook()
    ws = wb.active
    ws.title = "Kravsporing"

    headers = ["Søkeord", "Korttekst", "Funnet tekst", "Kravtype", "Treff %", "Dokument og side",
               "Valgt løsning", "Risiko", "Kommentar", "Kravverb", "Antall setninger", "Antall tegn",
               "NS-dokument", "NS-referanse"]

    ws.append(headers)
    ws.append([])
    ws.append(["Rapportoppsummering:"])
    ws.append([f"Totalt antall skannede dokumenter: {total_files_scanned}"])
    ws.append([f"Totalt antall unike krav funnet: {total_requirements_found}"])
    ws.append([])
    ws.append(["Fordeling av trefftyper:"])
    ws.append([f"  - Treff fra eksakte nøkkelord: {sum(keyword_counts.values())}"])
    ws.append([f"  - Treff fra eksakte synonymer: {sum(synonym_counts.values())}"])
    ws.append([f"  - Treff fra Maskinlæresøk (AI-søk): {sum(semantic_search_counts.values())}"])
    ws.append([])
    ws.append(headers)
    filter_row = ws.max_row
    ws.auto_filter.ref = f"A{filter_row}:{get_column_letter(len(headers))}{filter_row}"

    gray_fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
    white_fill = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
    bold_font = Font(bold=True)
    thin_border = Border(
        left=Side(style="thin", color="000000"),
        right=Side(style="thin", color="000000"),
        top=Side(style="thin", color="000000"),
        bottom=Side(style="thin", color="000000")
    )

    samlet_krav = [k for kravliste in all_requirements.values() for k in kravliste]
    samlet_krav.sort(key=lambda x: (x["keyword"].lower(), -x["score"]))

    start_row = ws.max_row + 1
    for r_idx, krav in enumerate(samlet_krav, start=start_row):
        nlp_v = krav.get("nlp_vurdering", {})
        kravverb = "Ja" if nlp_v.get("inneholder_kravsverb") else "Nei"
        ant_setn = nlp_v.get("antall_setninger", "")
        ant_tegn = nlp_v.get("lengde", "")

        ns_treff = krav.get("ns_treff", [])
        std_dok = ", ".join(sorted(set(t["standard"] for t in ns_treff))) if ns_treff else ""
        utdrag = []
        for t in ns_treff:
            preview = t['tekst'][:100].replace('\n',' ').strip()
            if len(t['tekst']) > 100: preview += "..."
            utdrag.append(f"{t['standard']} (s.{t['side']}): {preview} [Score: {t['score']:.1f}%]")
        ns_ref = "\n".join(utdrag) if utdrag else "Ingen relevante treff."

        row_data = [
            krav["keyword"],
            krav.get("short_text", ""),
            krav["text"],
            krav["kravtype"],
            krav["score"]/100.0,
            krav["ref"],
            "→ Beskriv valgt løsning her",
            "",
            "",
            kravverb,
            ant_setn,
            ant_tegn,
            std_dok,
            ns_ref
        ]
        ws.append(row_data)

        for c in range(1, len(headers)+1):
            cell = ws.cell(row=r_idx, column=c)
            cell.alignment = Alignment(wrap_text=True, vertical='top')
            cell.border = thin_border
            if c == 1:
                cell.fill = gray_fill; cell.font = bold_font
            elif c == 2:
                cell.fill = white_fill
            elif c == 4:
                kt = krav["kravtype"].lower()
                col = KRAVTYPE_COLORS.get(kt, KRAVTYPE_COLORS["generell"])
                cell.fill = PatternFill(start_color=col, end_color=col, fill_type="solid")
            elif c == 5:
                cell.number_format = "0%"

    for idx, header in enumerate(headers, start=1):
        width = 50 if header in ["Funnet tekst", "Valgt løsning", "Kommentar", "NS-referanse"] \
            else 40 if header == "Korttekst" \
            else 30 if header == "Dokument og side" else 20
        ws.column_dimensions[get_column_letter(idx)].width = width

    buf = BytesIO(); wb.save(buf); buf.seek(0)
    return buf

# === Excel-rapport (GRUPPERT) ===
def generate_total_excel_report_grouped(
    all_requirements: dict,
    keywords: list,
    total_files_scanned: int,
    total_requirements_found: int,
    keyword_counts: defaultdict,
    synonym_counts: defaultdict,
    semantic_search_counts: defaultdict,
    grouped_requirements: defaultdict
) -> BytesIO:
    wb = Workbook()
    ws_sum = wb.active
    ws_sum.title = "Sammendrag"

    # --- Sammendrag ---
    ws_sum.append(["Kravsporing – Sammendrag"])
    ws_sum.append([])
    ws_sum.append(["Totalt antall skannede dokumenter", total_files_scanned])
    ws_sum.append(["Totalt antall unike krav funnet", total_requirements_found])
    ws_sum.append([])
    ws_sum.append(["Trefftyper"])
    ws_sum.append(["Eksakte nøkkelord", sum(keyword_counts.values())])
    ws_sum.append(["Eksakte synonymer", sum(synonym_counts.values())])
    ws_sum.append(["Maskinlæresøk (AI)", sum(semantic_search_counts.values())])
    ws_sum.append([])

    # Fordeling per gruppe
    ws_sum.append(["Fordeling per gruppe"])
    ws_sum.append(["Gruppe", "Antall krav"])
    for g in sorted(grouped_requirements.keys(), key=lambda k: (k=="Uspesifisert", k)):
        ws_sum.append([g, len(grouped_requirements[g])])

    # Stretch litt kolonner
    for col in range(1, 3):
        ws_sum.column_dimensions[get_column_letter(col)].width = 40

    # --- Alle funn (med gruppe) ---
    ws_all = wb.create_sheet("Alle funn")
    headers = ["Søkeord", "Korttekst", "Funnet tekst", "Kravtype", "Treff %", "Dokument og side",
               "Valgt løsning", "Risiko", "Kommentar", "Kravverb", "Antall setninger", "Antall tegn",
               "NS-dokument", "NS-referanse", "Gruppe"]
    ws_all.append(headers)

    gray_fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
    bold_font = Font(bold=True)
    thin_border = Border(
        left=Side(style="thin", color="000000"),
        right=Side(style="thin", color="000000"),
        top=Side(style="thin", color="000000"),
        bottom=Side(style="thin", color="000000")
    )

    samlet_krav = [k for kravliste in all_requirements.values() for k in kravliste]
    # legg på gruppe (AI + fallback) hvis mangler
    for it in samlet_krav:
        if not it.get("gruppe"):
            best_fag, _ = classify_group_ai(it["text"])
            it["gruppe"] = best_fag or classify_group(it["text"])

    samlet_krav.sort(key=lambda x: (x.get("gruppe","Uspesifisert"), x["keyword"].lower(), -x["score"]))

    for krav in samlet_krav:
        nlp_v = krav.get("nlp_vurdering", {})
        kravverb = "Ja" if nlp_v.get("inneholder_kravsverb") else "Nei"
        ant_setn = nlp_v.get("antall_setninger", "")
        ant_tegn = nlp_v.get("lengde", "")

        ns_treff = krav.get("ns_treff", [])
        std_dok = ", ".join(sorted(set(t["standard"] for t in ns_treff))) if ns_treff else ""
        utdrag = []
        for t in ns_treff:
            preview = t['tekst'][:100].replace('\n',' ').strip()
            if len(t['tekst']) > 100: preview += "..."
            utdrag.append(f"{t['standard']} (s.{t['side']}): {preview} [Score: {t['score']:.1f}%]")
        ns_ref = "\n".join(utdrag) if utdrag else "Ingen relevante treff."

        ws_all.append([
            krav["keyword"],
            krav.get("short_text",""),
            krav["text"],
            krav["kravtype"],
            krav["score"]/100.0,
            krav["ref"],
            "→ Beskriv valgt løsning her",
            "",
            "",
            kravverb,
            ant_setn,
            ant_tegn,
            std_dok,
            ns_ref,
            krav.get("gruppe","Uspesifisert"),
        ])

    # format
    for r in ws_all.iter_rows(min_row=2, max_row=ws_all.max_row, min_col=1, max_col=len(headers)):
        for c in r:
            c.alignment = Alignment(wrap_text=True, vertical='top')
            c.border = thin_border
    # header styling
    for i, h in enumerate(headers, start=1):
        cell = ws_all.cell(row=1, column=i)
        cell.font = bold_font
        cell.fill = gray_fill
        ws_all.column_dimensions[get_column_letter(i)].width = (
            50 if h in ["Funnet tekst","Valgt løsning","Kommentar","NS-referanse"] else
            40 if h in ["Korttekst"] else
            30 if h in ["Dokument og side"] else
            22
        )
    ws_all.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{ws_all.max_row}"

    # --- Ett ark per gruppe ---
    for group_name in sorted({it.get("gruppe","Uspesifisert") for it in samlet_krav}, key=lambda k: (k=="Uspesifisert", k)):
        ws_g = wb.create_sheet(group_name[:31])
        ws_g.append(headers)
        group_rows = [it for it in samlet_krav if it.get("gruppe","Uspesifisert")==group_name]
        for krav in group_rows:
            nlp_v = krav.get("nlp_vurdering", {})
            kravverb = "Ja" if nlp_v.get("inneholder_kravsverb") else "Nei"
            ant_setn = nlp_v.get("antall_setninger", "")
            ant_tegn = nlp_v.get("lengde", "")
            ns_treff = krav.get("ns_treff", [])
            std_dok = ", ".join(sorted(set(t["standard"] for t in ns_treff))) if ns_treff else ""
            utdrag = []
            for t in ns_treff:
                preview = t['tekst'][:100].replace('\n',' ').strip()
                if len(t['tekst']) > 100: preview += "..."
                utdrag.append(f"{t['standard']} (s.{t['side']}): {preview} [Score: {t['score']:.1f}%]")
            ns_ref = "\n".join(utdrag) if utdrag else "Ingen relevante treff."

            ws_g.append([
                krav["keyword"], krav.get("short_text",""), krav["text"], krav["kravtype"],
                krav["score"]/100.0, krav["ref"], "→ Beskriv valgt løsning her", "", "",
                kravverb, ant_setn, ant_tegn, std_dok, ns_ref, group_name
            ])
        # format enkelt
        for r in ws_g.iter_rows(min_row=1, max_row=ws_g.max_row, min_col=1, max_col=len(headers)):
            for c in r:
                c.alignment = Alignment(wrap_text=True, vertical='top')
                c.border = thin_border
        for i, h in enumerate(headers, start=1):
            cell = ws_g.cell(row=1, column=i); cell.font = bold_font; cell.fill = gray_fill
            ws_g.column_dimensions[get_column_letter(i)].width = (
                50 if h in ["Funnet tekst","Valgt løsning","Kommentar","NS-referanse"] else
                40 if h in ["Korttekst"] else
                30 if h in ["Dokument og side"] else
                22
            )
        ws_g.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{ws_g.max_row}"

    buf = BytesIO()
    wb.save(buf); buf.seek(0)
    return buf

# === Fargekodet leverandør-rapport (Word) ===
def generate_colored_delivery_report(grouped_requirements: defaultdict) -> io.BytesIO:
    doc = Document()
    doc.add_heading("Leverandørleveranser", 0)
    doc.add_paragraph(
        "Denne rapporten grupperer funne krav basert på potensielt ansvarlig fagområde/leverandør. "
        "Verifiser grupperingen før bruk."
    )

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Gruppe", "Krav", "Type", "Referanse"

    # Header shading
    for c in hdr:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), "C0C0C0")
        c._tc.get_or_add_tcPr().append(shd)

    def filtrer_unike(kravliste: list, terskel=95):
        unike = []
        for item in kravliste:
            if not any(fuzz_ratio(item['text'].lower(), ex['text'].lower()) >= terskel for ex in unike):
                unike.append(item)
        return unike

    for grp in sorted(grouped_requirements.keys(), key=lambda g: (g == "Uspesifisert", g)):
        for krav in sorted(filtrer_unike(grouped_requirements[grp]), key=lambda x: x['score'], reverse=True):
            row = table.add_row().cells
            row[0].text, row[1].text, row[2].text, row[3].text = grp, krav["text"], krav["kravtype"], krav["ref"]
            if grp in GROUP_COLORS:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), GROUP_COLORS[grp])
                for c in row:
                    c._tc.get_or_add_tcPr().append(shd)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# === Grafisk oversikt (Word med bilder) ===
def generate_graphical_report_docx(
    total_requirements_found: int,
    keyword_counts: defaultdict,
    synonym_counts: defaultdict,
    semantic_search_counts: defaultdict,
    grouped_requirements: defaultdict,
    all_requirements: dict
) -> io.BytesIO:
    """
    Lager en Word-rapport med tre grafer (pie + 2 bar) uten å bruke interaktiv GUI-backend.
    Sikker i Celery/threads: tvinger Matplotlib til 'Agg' og slår av interaktiv modus.
    """
    import matplotlib
    try:
        matplotlib.use("Agg", force=True)
    except Exception:
        pass
    import matplotlib.pyplot as plt
    plt.ioff()

    doc = Document()
    doc.add_heading('Grafisk Kravsporingsoversikt', 0)
    doc.add_paragraph(f"Rapport generert: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Totalt antall unike krav identifisert: {total_requirements_found}")
    doc.add_paragraph("")

    os.makedirs(TEMP_ROOT, exist_ok=True)
    ts = int(time.time())

    # --- Fordeling av trefftyper (pie) ---
    try:
        labels_pie = ['Eksakte Nøkkelord', 'Eksakte Synonymer', 'Maskinlæresøk (AI-søk)']
        sizes_pie = [sum(keyword_counts.values()), sum(synonym_counts.values()), sum(semantic_search_counts.values())]
        if sum(sizes_pie) > 0:
            fig1, ax1 = plt.subplots(figsize=(8, 8))
            ax1.pie(sizes_pie, labels=labels_pie, autopct='%1.1f%%', startangle=90)
            ax1.axis('equal')
            ax1.set_title('Fordeling av Trefftyper')
            pie_path = TEMP_ROOT / f"trefftyper_pie_{ts}.png"
            plt.savefig(pie_path, bbox_inches='tight')
            plt.close(fig1)
            doc.add_heading('Fordeling av Trefftyper', level=1)
            doc.add_picture(str(pie_path), width=Inches(6))
            doc.add_paragraph("")
    except Exception as e:
        logging.warning(f"Kunne ikke lage pie-chart: {e}", exc_info=True)

    # --- Fordeling av kravtyper (bar) ---
    try:
        kravtype_counts = defaultdict(int)
        for krav_list in all_requirements.values():
            for item in krav_list:
                kravtype_counts[item['kravtype']] += 1
        if kravtype_counts:
            labels_bar = list(kravtype_counts.keys())
            counts_bar = list(kravtype_counts.values())
            fig2, ax2 = plt.subplots(figsize=(10, 6))
            ax2.bar(labels_bar, counts_bar)
            ax2.set_ylabel('Antall krav')
            ax2.set_title('Fordeling av Kravtyper')
            ax2.tick_params(axis='x', rotation=45)
            bar_path = TEMP_ROOT / f"kravtyper_bar_{ts}.png"
            plt.tight_layout()
            plt.savefig(bar_path, bbox_inches='tight')
            plt.close(fig2)
            doc.add_heading('Fordeling av Kravtyper', level=1)
            doc.add_picture(str(bar_path), width=Inches(6))
            doc.add_paragraph("")
    except Exception as e:
        logging.warning(f"Kunne ikke lage bar-chart (kravtyper): {e}", exc_info=True)

    # --- Fordeling per gruppe (bar) ---
    try:
        if grouped_requirements:
            order = sorted(grouped_requirements.keys(), key=lambda g: (g == "Uspesifisert", g))
            labels_group = [g.capitalize() for g in order]
            counts_group = [len(grouped_requirements[g]) for g in order]
            fig3, ax3 = plt.subplots(figsize=(10, 6))
            ax3.bar(labels_group, counts_group)
            ax3.set_ylabel('Antall krav')
            ax3.set_title('Fordeling av Fagområder / Leverandørgrupper')
            ax3.tick_params(axis='x', rotation=45)
            grp_path = TEMP_ROOT / f"grupper_bar_{ts}.png"
            plt.tight_layout()
            plt.savefig(grp_path, bbox_inches='tight')
            plt.close(fig3)
            doc.add_heading('Fordeling av Fagområder / Leverandørgrupper', level=1)
            doc.add_picture(str(grp_path), width=Inches(6))
            doc.add_paragraph("")
    except Exception as e:
        logging.warning(f"Kunne ikke lage bar-chart (grupper): {e}", exc_info=True)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# === Oppsummering ===
def generate_summary_report(
    all_requirements: dict,
    keywords: list,
    total_files_scanned: int,
    total_requirements_found: int,
    keyword_counts: defaultdict,
    synonym_counts: defaultdict,
    top_score_krav: dict,
    grouped_requirements: defaultdict,
    semantic_search_counts: defaultdict,
    suggested_keywords: list = None
) -> io.BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    doc.add_heading('Oppsummering av utført kravsporing', 0)
    doc.add_paragraph(f"Generert: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    if keywords:
        doc.add_paragraph(f"Nøkkelord brukt i søket: {', '.join(keywords)}")
    else:
        doc.add_paragraph("Analysemodus: Bare AI-modell (ingen nøkkelord).")
    doc.add_paragraph("")

    doc.add_heading('Søke- og funnstatistikk', 1)
    doc.add_paragraph(f"Totalt antall dokumenter skannet: {total_files_scanned}")
    doc.add_paragraph(f"Totalt antall unike krav identifisert: {total_requirements_found}")
    doc.add_paragraph("")

    doc.add_heading('Funn per kategori', 2)
    if keyword_counts:
        doc.add_paragraph("Eksakte nøkkelord:")
        for kw, c in sorted(keyword_counts.items(), key=lambda x: x[1], reverse=True):
            doc.add_paragraph(f"- '{kw}': {c} treff")
    else:
        doc.add_paragraph("Ingen eksakte nøkkelordstreff.")
    doc.add_paragraph("")
    if synonym_counts:
        doc.add_paragraph("Eksakte synonymer:")
        for syn, c in sorted(synonym_counts.items(), key=lambda x: x[1], reverse=True):
            doc.add_paragraph(f"- '{syn}': {c} treff")
    else:
        doc.add_paragraph("Ingen eksakte synonymtreff.")
    doc.add_paragraph("")
    if semantic_search_counts:
        doc.add_paragraph("Maskinlæresøk (AI-søk):")
        for term, c in sorted(semantic_search_counts.items(), key=lambda x: x[1], reverse=True):
            doc.add_paragraph(f"- '{term}': {c} treff")
    else:
        doc.add_paragraph("Ingen AI-treff.")
    doc.add_paragraph("")

    doc.add_heading('Høyest scorede krav', 1)
    if top_score_krav:
        doc.add_paragraph("Høyeste score i hele kjøringen:")
        doc.add_paragraph(
            f"Score: {top_score_krav['score']:.1f}% | "
            f"Søkeord: '{top_score_krav['keyword']}' | "
            f"Tekst: '{top_score_krav['text']}' | "
            f"Ref: {top_score_krav['ref']}"
        )
    else:
        doc.add_paragraph("Ingen krav funnet.")
    doc.add_paragraph("")

    doc.add_heading('Oppsummering etter fag/gruppe', 1)
    if grouped_requirements:
        for group_name in sorted(grouped_requirements.keys(), key=lambda g: (g == "Uspesifisert", g)):
            def filtrer_unike(lst, t=95):
                unike = []
                for k in lst:
                    if not any(fuzz_ratio(k['text'].lower(), e['text'].lower()) >= t for e in unike):
                        unike.append(k)
                return unike
            uniq = filtrer_unike(grouped_requirements[group_name])
            doc.add_heading(f"{group_name.capitalize()} ({len(uniq)} krav)", 2)
            for krav_item in sorted(uniq, key=lambda x: x['score'], reverse=True)[:3]:
                doc.add_paragraph(f"• {krav_item['text'][:150]}... (Score: {krav_item['score']:.1f}%)")
            doc.add_paragraph("")
    else:
        doc.add_paragraph("Ingen grupperte krav identifisert.")
    doc.add_paragraph("")

    doc.add_heading('Foreslåtte nye nøkkelord', 1)
    if suggested_keywords:
        for kw in suggested_keywords:
            doc.add_paragraph(f"- {kw}")
    else:
        doc.add_paragraph("Ingen forslag tilgjengelig.")

    doc.add_heading('Konklusjon/anbefalinger', 1)
    for _ in range(5):
        doc.add_paragraph("____________________________________________________________________")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out

def _split_atomic_requirements(text: str) -> list[str]:
    """
    Splitt en normativ blokk i 'atomiske' krav (én setning/én påstand).
    Tåler tett PDF-tekst, semikolon, punktum, bullets og kolon-lister.
    """
    if not text:
        return []

    # Normaliser noen typiske PDF-artefakter
    t = text.replace("•", " • ").replace("·", " • ")
    # Grov splitting: punktum/semikolon/linjeskift/bullets/kolon som ofte introduserer ny påstand
    parts = re.split(r'(?<=[.;])\s+|\n+|\s*•\s*|\s*:\s*', t)

    # Start-nøkkelord som ofte markerer ny påstand inne i samme linje
    re_lead = re.compile(
        r'\b('
        r'SFP|VAV|CO2|CO₂|varmegjenvinning|operativ temperatur|romtemperatur|tilluftstemperatur|'
        r'DUT|dimensjonerende utetemperatur|utetemperatur|'
        r'luftmengde|sekvensregulering|tilluft|avtrekk|'
        r'undertrykk|overtrykk|trykkdifferanse|differansetrykk|trykksetting'
        r')\b', re.I
    )

    out = []
    for p in parts:
        p = p.strip()
        if not p:
            continue

        # Hvis flere “ledende” nøkkelord finnes i samme bit, splitt på dem
        chunks = []
        start = 0
        for m in re_lead.finditer(p):
            if m.start() > start:
                chunks.append(p[start:m.start()].strip())
            start = m.start()
        chunks.append(p[start:].strip())

        for c in chunks:
            if len(c) < 12:
                continue
            cl = c.lower()
            # behold bare normative setninger (kravsverb / enheter / domeneord)
            if any(v in cl for v in GLOBAL_OBLIGATION_VERBS) or UNITS_REGEX.search(cl) or re_lead.search(c):
                # sikre at setningen avsluttes pent
                if not c.endswith(('.', ';')):
                    c = c + '.'
                out.append(c)

    # fjern trivielle dubletter
    out = list(dict.fromkeys(out))
    return out

def _has_verb(text: str) -> bool:
    tl = (text or "").lower()
    if nlp:
        try:
            return any(t.pos_ == "VERB" for t in nlp(text))
        except Exception:
            pass
    return bool(re.search(r"\b(skal|må|kan|er|være|forutsettes|leveres|etableres|plasseres|utstyres|tilpasses|legges|dimensjoneres)\b", tl))

def _looks_incomplete(s: str) -> bool:
    # Slutter på preposisjon/komma/ingen verb → sannsynlig fortsettelse
    tl = (s or "").strip().lower()
    if not tl:
        return True
    if tl.endswith((",", ";", ":", " og", " samt", " med", " for", " til", " på", " av", " i", " om")):
        return True
    if not _has_verb(tl) and len(tl) < 80:
        return True
    return False

def _starts_as_continuation(s: str) -> bool:
    # Starter med liten bokstav/konjunksjon/kommategn → sannsynlig fortsettelse
    tl = (s or "").lstrip()
    if not tl:
        return False
    return bool(re.match(r"^(og|samt|samtidig|slik at|der|som)\b", tl.lower())) or bool(re.match(r"^[a-zæøå]", tl))

def _enrich_clause(clauses: list[str], idx: int, profile: dict) -> tuple[str, int]:
    """
    Bygg 'rik' kravtekst rundt clauses[idx]:
      - prepend forrige dersom den bærer subjekt/omfang
      - append én eller flere neste dersom basis er ufullstendig / neste ser ut som fortsettelse
      - append neste som har tall/enheter hvis basis mangler tall
    Returnerer (beriket_tekst, antall_forbrukte_clauses).
    """
    used = 1
    base = clauses[idx].strip()

    # 1) Prepend forrige (subjekt/omfang) hvis basis mangler det
    if (_looks_incomplete(base) or not _has_verb(base)) and idx - 1 >= 0:
        prev = clauses[idx - 1].strip()
        # prioritér kort forrige uten tall men med domenetermer/kravsverb
        prev_ok = (_has_verb(prev) or any(t in prev.lower() for t in profile.get("terms", []))) and not _contains_numbers_units_profile(prev, profile)
        if prev_ok and len(prev) < 160:
            base = prev.rstrip(" .;") + ". " + base

    # 2) Append “fortsettelser” bakover-til-framover så lenge det trengs
    j = idx + 1
    while j < len(clauses) and (_looks_incomplete(base) or _starts_as_continuation(clauses[j])):
        nxt = clauses[j].strip()
        if not nxt:
            break
        base = base.rstrip(" .;") + ". " + nxt
        used += 1
        j += 1

    # 3) Hvis fortsatt ingen tall/enheter → append første neste som har tall/enheter
    if not _contains_numbers_units_profile(base, profile) and j < len(clauses):
        nxt = clauses[j].strip()
        if _contains_numbers_units_profile(nxt, profile):
            base = base.rstrip(" .;") + ". " + nxt
            used += 1
            j += 1

    # Rydd opp tegnsetting
    base = re.sub(r"\s+", " ", base).strip()
    base = base.replace(", .", ".").replace(" ,", ",")
    if not base.endswith((".", ";")):
        base += "."

    return base, used

def extract_requirements(
    text: str,
    base_keywords: list,
    file_name: str,
    min_score: float,
    file_type: str,
    ns_standard_selection: str = "Ingen",
    mode: str = "keywords_ai",
    fokus_text: str = "",
    use_fokus_prefilter: bool = True,
    fokus_threshold: float = 0.60,
    selected_groups: list | None = None
) -> list:
    # --- last synonymer ---
    current_synonyms = {}
    if SYNONYM_PATH.exists():
        try:
            with open(SYNONYM_PATH, "r", encoding="utf-8") as f:
                raw = json.load(f) or {}
            current_synonyms = {k.lower(): [s.lower() for s in (v or [])] for k, v in raw.items()}
        except json.JSONDecodeError:
            pass

    # --- segmenter tekst til setninger + sidepeker ---
    sentences, page_for_sent = [], []
    parts = re.split(r'\[\[SIDE\s+(\d+)\]\]', text)
    if len(parts) >= 3:
        # parts = [ev. tekst før første markør, page1, text1, page2, text2, ...]
        it = iter(parts[1:])
        for page_no, page_text in zip(it, it):
            page_no = str(page_no).strip()
            if nlp:
                doc = nlp(page_text)
                for s in doc.sents:
                    st = s.text.strip()
                    if st:
                        sentences.append(st)
                        page_for_sent.append(page_no)
            else:
                for st in re.split(r'(?<=[.!?])\s+', page_text):
                    st = st.strip()
                    if st:
                        sentences.append(st)
                        page_for_sent.append(page_no)
    else:
        # Fallback hvis ingen markører — burde være sjeldent
        current_page = "Ukjent"
        if nlp:
            doc = nlp(text)
            for s in doc.sents:
                st = s.text.strip()
                if st:
                    sentences.append(st)
                    page_for_sent.append(current_page)
        else:
            for st in re.split(r'(?<=[.!?])\s+', text):
                st = st.strip()
                if st:
                    sentences.append(st)
                    page_for_sent.append(current_page)

    if not sentences:
        return []

    # Velg aktiv domenprofil fra fokus-tekst
    active_domain = _choose_domain_from_inputs(fokus_text, selected_groups)
    profile = DOMAIN_PROFILES.get(active_domain, DOMAIN_PROFILES["generic"])

    # --- FOKUS-PREFILTER: bruk fokus-teksten allerede under skanningen ---
    # if use_fokus_prefilter and fokus_text and semantic_model and sentences:
    #     try:
    #         f_emb = semantic_model.encode(fokus_text, convert_to_tensor=True)
    #         s_embs = semantic_model.encode(sentences, convert_to_tensor=True, show_progress_bar=False)
    #         sims = util.pytorch_cos_sim(f_emb, s_embs)[0]
    #         kept_sentences, kept_pages = [], []
    #         for i, _ in enumerate(sentences):
    #             if sims[i].item() >= fokus_threshold:
    #                 kept_sentences.append(sentences[i])
    #                 kept_pages.append(page_for_sent[i])
    #         if kept_sentences:
    #             sentences = kept_sentences
    #             page_for_sent = kept_pages
    #             logging.info(f"[FOKUS-PREFILTER] Beholdt {len(sentences)} setninger for fokus '{fokus_text[:60]}...'")
    #         else:
    #             logging.info("[FOKUS-PREFILTER] Ingen setninger passerte terskel – beholder originalt av sikkerhet.")
    #     except Exception as e:
    #         logging.warning(f"[FOKUS-PREFILTER] Feilet, hopper over: {e}", exc_info=False)

    results, seen = [], set()

    # --- forbered nøkkelord/embedding ---
    kw_embeds = None
    base_lc = [k.lower() for k in (base_keywords or [])]
    if mode in {"keywords_ai"} and semantic_model and base_lc:
        try:
            kw_embeds = semantic_model.encode(base_keywords, convert_to_tensor=True)
        except Exception:
            kw_embeds = None

    def link_ns(context_text: str):
        if not semantic_model or not global_ns_data:
            return []
        try:
            q_emb = semantic_model.encode(context_text, convert_to_tensor=True)
            hits = []
            # respekter eventuelt UI-valg: ns_standard_selection
            active_standards = [
                s for s, cfg in PDF_STANDARDER.items()
                if cfg.get("aktiv", False)
            ]
            # Respekter UI-valget
            if ns_standard_selection and ns_standard_selection != "Ingen":
                active_standards = [ns_standard_selection] if ns_standard_selection in PDF_STANDARDER else []
            else:
                # "Ingen" valgt → slå av NS-søk
                return []
            if not active_standards:
                return []
            for std in active_standards:
                entries = global_ns_data.get(std, [])
                if not entries:
                    continue
                corpus_embs = [e["embedding"] for e in entries]
                # top_k lite for ytelse
                matches = util.semantic_search(q_emb, corpus_embs, top_k=3)[0]
                for m in matches:
                    e = entries[m["corpus_id"]]
                    hits.append({
                        "standard": std,
                        "side": e["side"],
                        "tekst": e["tekst"],
                        "score": float(m["score"] * 100.0),
                    })
            # sorter og klipp
            hits.sort(key=lambda x: x["score"], reverse=True)
            return hits[:5]
        except Exception:
            return []

    kw_regexes = defaultdict(list)
    for kw in base_lc:
        kw_regexes[kw].append(re.compile(r'\b' + re.escape(kw) + r'\b', flags=re.I))
        if current_synonyms.get(kw):
            for syn in current_synonyms[kw]:
                kw_regexes[kw].append(re.compile(r'\b' + re.escape(syn.lower()) + r'\b', flags=re.I))

    # --- gå gjennom setninger og bygg normative blokker ---
    i = 0
    while i < len(sentences):
        sent = sentences[i]
        page = page_for_sent[i]
        sent_lc = sent.lower()

        key_hit, score = None, 0.0

        # Eksakt/synonymt nøkkelord
        if mode in {"keywords", "keywords_ai"} and base_lc:
            for kw in base_lc:
                if kw_regexes[kw][0].search(sent_lc):
                    key_hit, score = kw, 100.0
                    break
            if not key_hit:
                for kw in base_lc:
                    for rx in kw_regexes[kw][1:]:
                        if rx.search(sent_lc):
                            key_hit, score = kw, 95.0
                            break
                    if key_hit: break

        # Semantisk nøkkelord – kun i keywords_ai
        if not key_hit and mode == "keywords_ai" and semantic_model and kw_embeds is not None and base_lc:
            try:
                s_emb = semantic_model.encode(sent, convert_to_tensor=True)
                matches = util.semantic_search(s_emb, kw_embeds, top_k=1)
                if matches and matches[0] and (matches[0][0]['score'] * 100 >= min_score):
                    best = matches[0][0]
                    key_hit = base_keywords[best['corpus_id']]
                    score = best['score'] * 100
            except Exception:
                pass

        # Ren AI-modus
        if mode == "ai" and not key_hit:
            if _is_normative_sentence_profile(sent, profile):
                key_hit = "(AI)"
                score = 88.0 if any(v in sent_lc for v in GLOBAL_OBLIGATION_VERBS) else 85.0


        # Hvis ingen treff – gå videre
        if not key_hit:
            i += 1
            continue

        # Bygg normativ blokk fra denne ankersetningen
        block_text, last_idx = _build_normative_block(sentences, i, profile)
        clauses = _split_atomic_requirements(block_text)
        made_any = False
        j = 0
        while j < len(clauses):
            rich_text, step = _enrich_clause(clauses, j, profile)

            if not is_valid_requirement(rich_text):
                j += step
                continue

            context = rich_text
            loc = (context, page)
            if loc in seen:
                j += step
                continue
            seen.add(loc)

            d = nlp(context) if nlp else None
            is_pressure = re.search(r"\b(undertrykk|overtrykk|trykk(?:differanse|setting|settes|satt))\b", context, re.I)
            is_temp     = re.search(r"\b(DUT|dimensjonerende|dimensjonerende utetemperatur|tilluftstemperatur|operativ temperatur|romtemperatur|utetemperatur)\b", context, re.I)
            boost = (3.0 if _contains_numbers_units_profile(context, profile) else 0.0) \
                    + (4.0 if is_pressure else 0.0) \
                    + (2.0 if is_temp else 0.0)

            results.append({
                "keyword": key_hit,
                "text": context,                              # ← nå fyldig, komplett setning/sett
                "kravtype": classify_type(context),
                "score": min(100.0, (score or 0.0) + boost),
                "ref": f"{file_name} / Side {page}",
                "short_text": _generate_short_text(context),
                "anchor": sent,
                "nlp_vurdering": {
                    "inneholder_kravsverb": any(tok.lemma_.lower() in GLOBAL_OBLIGATION_VERBS for tok in d if tok.pos_ == "VERB") if d else False,
                    "antall_setninger": len(list(d.sents)) if d else 1,
                    "lengde": len(context)
                },
                "ns_treff": link_ns(context)
            })
            made_any = True
            j += step

        # fallback hvis blokken ikke ga noe
        if not made_any and is_valid_requirement(block_text):
            context = block_text
            loc = (context, page)
            if loc not in seen:
                seen.add(loc)
                d = nlp(context) if nlp else None
                results.append({
                    "keyword": key_hit,
                    "text": context,
                    "kravtype": classify_type(context),
                    "score": score or 0.0,
                    "ref": f"{file_name} / Side {page}",
                    "short_text": _generate_short_text(context),
                    "anchor": sent,
                    "nlp_vurdering": {
                        "inneholder_kravsverb": any(tok.lemma_.lower() in GLOBAL_OBLIGATION_VERBS for tok in d if tok.pos_ == "VERB") if d else False,
                        "antall_setninger": len(list(d.sents)) if d else 1,
                        "lengde": len(context)
                    },
                    "ns_treff": link_ns(context)
                })
            made_any = True

        # Fallback hvis ingenting ble laget
        if not made_any and is_valid_requirement(block_text):
            context = block_text
            loc = (context, page)
            if loc not in seen:
                seen.add(loc)
                d = nlp(context) if nlp else None
                results.append({
                    "keyword": key_hit,
                    "text": context,
                    "kravtype": classify_type(context),
                    "score": score or 0.0,
                    "ref": f"{file_name} / Side {page}",
                    "short_text": _generate_short_text(context),
                    "anchor": sent,
                    "nlp_vurdering": {
                        "inneholder_kravsverb": any(tok.lemma_.lower() in GLOBAL_OBLIGATION_VERBS for tok in d if tok.pos_ == "VERB") if d else False,
                        "antall_setninger": len(list(d.sents)) if d else 1,
                        "lengde": len(context)
                    },
                    "ns_treff": link_ns(context)
                })
            made_any = True

        # Fallback: hvis ingen gyldige delkrav ble laget, ta hele blokken som før
        if not made_any and is_valid_requirement(block_text):
            context = block_text
            loc = (context, page)
            if loc not in seen:
                seen.add(loc)
                d = nlp(sent) if nlp else None
                results.append({
                    "keyword": key_hit,
                    "text": context,
                    "kravtype": classify_type(context),
                    "score": score or 0.0,
                    "ref": f"{file_name} / Side {page}",
                    "short_text": _generate_short_text(context),
                    "anchor": sent,
                    "nlp_vurdering": {
                        "inneholder_kravsverb": any(tok.lemma_.lower() in GLOBAL_OBLIGATION_VERBS for tok in d if tok.pos_ == "VERB") if d else False,
                        "antall_setninger": len(list(d.sents)) if d else 1,
                        "lengde": len(sent)
                    },
                    "ns_treff": link_ns(context)
                })

        # hopp over setninger som ble slukt i blokken
        i = last_idx + 1

    results.sort(key=lambda x: x['score'], reverse=True)
    return results

BASE_GROUPS = ["elektro", "ventilasjon", "rørlegger", "byggautomasjon", "kulde", "totalentreprenør", "byggherre", "økonomi"]

# === Fag-normalisering (NYTT) ===
CANON = {
    "alle fag": "__ALL__",
    "elektro": "elektro",
    "ventilasjon": "ventilasjon",
    "rør": "rørlegger",
    "rørlegger": "rørlegger",
    "byggautomasjon": "byggautomasjon",
    "totalentreprenør": "totalentreprenør",
    "byggherre": "byggherre",
    "økonomi": "økonomi",
    "kulde": "kulde",
    "uspesifisert": "Uspesifisert",
}

def _canon_group_name(name: str) -> str:
    if not name:
        return "Uspesifisert"
    key = name.strip().lower()
    # eksakte alias
    alias = {
        "alle fag": "__ALL__",
        "rør": "rørlegger",
        "uspesifisert": "Uspesifisert",
    }
    if key in alias:
        return alias[key]
    # ellers behold som lower
    return key

def _base_group(name: str) -> str:
    """Returner basisfaget (elektro/ventilasjon/...), selv om teksten er 'ventilasjon, prosjektering' osv."""
    if not name:
        return "Uspesifisert"
    s = name.strip().lower()
    for g in BASE_GROUPS:
        if g in s:
            return g
    return _canon_group_name(name)

def _normalize_selected_groups(selected_groups):
    if not selected_groups:
        return []
    norm = {_canon_group_name(g) for g in selected_groups if isinstance(g, str) and g.strip()}
    # Hvis "Alle fag" er valgt SAMMEN med andre → dropp "__ALL__"
    if "__ALL__" in norm and len(norm) > 1:
        norm.discard("__ALL__")
    # Kun "__ALL__" igjen → betyr ingen filtrering (behold)
    if not norm:
        return []
    return sorted(norm)

def _process_single_document(self_task, file_path: Path, keywords: list, min_score: float,
                             ns_standard_selection: str, mode: str, fokusomraade: str, selected_groups: list | None):
    """
    Forbedret funksjon som leser filer, inkludert vedlegg i .msg-filer,
    og sørger for at sidetall fra PDF blir med videre.
    Returnerer en liste av (krav, feil) for hver tekstblokk som blir funnet.
    """
    
    results = []
    
    def process_text_content(text_content: str, source_name: str):
        """Hjelpefunksjon for å kjøre kravuthenting på en tekstblokk."""
        file_reqs, file_errs = [], []
        try:
            if text_content:
                cleaned = clean_text(text_content)
                file_reqs = extract_requirements(
                    cleaned,
                    keywords,
                    source_name,
                    min_score,
                    source_name.split('.')[-1], # Bruk filtype fra kildenavnet
                    ns_standard_selection,
                    mode,
                    fokus_text=fokusomraade,
                    use_fokus_prefilter=True,
                    fokus_threshold=0.60,
                    selected_groups=selected_groups
                )
        except Exception as e:
            logging.error(f"FEIL under prosessering av innhold fra {source_name}: {e}", exc_info=True)
            file_errs.append(f"Alvorlig feil under prosessering av innhold fra '{source_name}': {e}")
        
        return file_reqs, file_errs

    filename = file_path.name
    fext = file_path.suffix.lower()

    if fext == ".msg":
        # Håndter e-post og vedlegg separat
        msg = extract_msg.Message(str(file_path))
        
        # 1. Prosesser selve e-postteksten
        email_body = msg.body
        if email_body:
            reqs, errs = process_text_content(email_body, f"{filename} (E-post)")
            results.append((reqs, errs))
        
        # 2. Gå gjennom og prosesser alle vedlegg
        for att in msg.attachments:
            if att.type == "data":
                att_filename = att.long_filename or att.short_filename
                att_ext = Path(att_filename).suffix.lower()
                att_content = ""
                
                try:
                    if att_ext == ".pdf":
                        pdf_doc = fitz.open(stream=att.data, filetype="pdf")
                        att_content = "\n".join([f"[[SIDE {i+1}]]\n{page.get_text('text')}" for i, page in enumerate(pdf_doc)])
                        pdf_doc.close()
                    elif att_ext == ".docx":
                        doc = Document(io.BytesIO(att.data))
                        att_content = "\n".join(p.text for p in doc.paragraphs)
                    elif att_ext == ".doc":
                        # Skriv vedlegget midlertidig til disk og konverter til docx
                        att_tmp_dir = TEMP_ROOT / f"doc_att_{secrets.token_hex(4)}"
                        att_tmp_dir.mkdir(parents=True, exist_ok=True)
                        try:
                            src_path = att_tmp_dir / (Path(att.long_filename or att.short_filename).name)
                            src_path.write_bytes(att.data)
                            docx_path = _convert_doc_to_docx(src_path, att_tmp_dir)
                            if docx_path and docx_path.exists():
                                att_content = _extract_text_from_docx_bytes(docx_path.read_bytes())
                        finally:
                            try:
                                shutil.rmtree(att_tmp_dir, ignore_errors=True)
                            except Exception:
                                pass
                    elif att_ext in [".txt", ".csv"]:
                        att_content = att.data.decode('utf-8', errors='ignore')
                    
                    if att_content:
                        reqs, errs = process_text_content(att_content, f"{filename} -> {att_filename}")
                        results.append((reqs, errs))

                except Exception as e:
                    logging.warning(f"Kunne ikke prosessere vedlegg {att_filename}: {e}")
                    results.append(([], [f"Kunne ikke lese vedlegg '{att_filename}' fra '{filename}'."]))

    else:
        # Håndter enkeltstående filer (PDF, DOCX, TXT etc.)
        extracted_text = ""
        try:
            if fext == ".pdf":
                # KORREKT HÅNDTERING: Inkluderer sidetall-markører
                with fitz.open(file_path) as doc:
                    extracted_text = "\n".join([f"[[SIDE {i+1}]]\n{page.get_text('text')}" for i, page in enumerate(doc)])
            elif fext == ".docx":
                doc = Document(file_path)
                extracted_text = "\n".join(p.text for p in doc.paragraphs)
            elif fext == ".doc":
                # Konverter lokalt .doc → .docx med LibreOffice
                conv_dir = TEMP_ROOT / f"doc_{secrets.token_hex(4)}"
                try:
                    docx_path = _convert_doc_to_docx(file_path, conv_dir)
                    if docx_path and docx_path.exists():
                        extracted_text = _extract_text_from_docx_bytes(docx_path.read_bytes())
                    else:
                        raise RuntimeError("Konvertering ga ingen .docx-utfil.")
                finally:
                    try:
                        shutil.rmtree(conv_dir, ignore_errors=True)
                    except Exception:
                        pass
            elif fext == '.txt':
                extracted_text = file_path.read_text(encoding="utf-8", errors='ignore')
            elif fext == ".xlsx":
                wb = load_workbook(file_path, read_only=True, data_only=True)
                extracted_text = "\n".join([str(c.value) for ws in wb.worksheets for row in ws.iter_rows() for c in row if c.value is not None])

            reqs, errs = process_text_content(extracted_text, filename)
            results.append((reqs, errs))

        except Exception as e:
            logging.error(f"FEIL ved lesing av fil {filename}: {e}", exc_info=True)
            results.append(([], [f"Kritisk feil ved lesing av fil '{filename}': {e}"]))

    # Slå sammen resultater fra alle tekstblokker (hoveddokument + vedlegg)
    final_reqs = [req for res_tuple in results for req in res_tuple[0]]
    final_errs = [err for res_tuple in results for err in res_tuple[1]]
    
    return final_reqs, final_errs

def _normalize_keywords(keywords):
    if not keywords:
        return []
    if isinstance(keywords, str):
        parts = re.split(r'[,\n;]+', keywords)
    elif isinstance(keywords, list):
        parts = []
        for k in keywords:
            if isinstance(k, str):
                parts.extend(re.split(r'[,\n;]+', k))
    else:
        parts = [str(keywords)]
    return [p.strip() for p in parts if p and p.strip()]

def _choose_domain_from_inputs(fokus_text: str = "", selected_groups: list | None = None) -> str:
    ft = (fokus_text or "").lower()
    # 1) Fokus-tekst vinner
    for key, prof in DOMAIN_PROFILES.items():
        if key == "generic": continue
        if any(a in ft for a in prof["aliases"]):
            return key
    # 2) Ellers fra valgte fag
    if selected_groups:
        s = " ".join([str(g).lower() for g in selected_groups])
        for key in DOMAIN_PROFILES.keys():
            if key != "generic" and key in s:
                return key
        if "rør" in s: return "rørlegger"
    return "generic"

def _profile_contains_cues(text: str, profile: dict) -> bool:
    tl = text.lower()
    return any(t in tl for t in profile["terms"]) or bool(profile["units_re"].search(tl))


def _normalize_for_fuzzy(s: str) -> str:
    # gjør tekster sammenlignbare (enhet- og symbol-normalisering)
    s = s.lower()
    s = s.replace("co₂", "co2").replace("°c", "c")
    s = s.replace(" m3 ", " m³ ").replace("m3/s", "m³/s").replace("m3/m2/t","m³/m²/t")
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"[^\wæøå0-9/ ]+", " ", s)
    s = re.sub(r"\bside\s*\d+\b", "", s, flags=re.I)
    return s.strip()

def _parse_ref(ref: str):
    m = re.match(r'^(.*?)\s*/\s*Side\s+(\d+)$', ref.strip(), flags=re.I)
    if m:
        try:
            return m.group(1), int(m.group(2))
        except Exception:
            return m.group(1), m.group(2)
    return ref, None

def _merge_refs(items):
    by_file = defaultdict(set)
    for it in items:
        fn, pg = _parse_ref(it.get("ref",""))
        if fn:
            if isinstance(pg, int): by_file[fn].add(pg)
            else: by_file[fn].add(pg if pg is not None else "")
    parts = []
    for fn in sorted(by_file.keys()):
        pages = sorted([p for p in by_file[fn] if isinstance(p, int)])
        others = [p for p in by_file[fn] if not isinstance(p, int) and p is not None]
        if pages:
            parts.append(f"{fn} / Side {', '.join(str(p) for p in pages)}")
        elif others:
            parts.append(f"{fn} / {', '.join(map(str, others))}")
        else:
            parts.append(fn)
    return "; ".join(parts)

def _norm_source(req):
    # Anker (selve treff-setningen) er beste basis for dedup; fall back til hele teksten
    base = req.get("anchor") or req.get("text","")
    return _normalize_for_fuzzy(base)

def _contains_domain_cues_profile(s: str, profile: dict) -> bool:
    tl = s.lower()
    return any(t in tl for t in profile.get("terms", [])) or bool(profile["units_re"].search(tl))

def _contains_numbers_units_profile(s: str, profile: dict) -> bool:
    return bool(profile["units_re"].search(s.lower()))

def _is_normative_sentence_profile(s: str, profile: dict | None = None) -> bool:
    """Normativ setning = kravsverb eller tall/enheter eller fag-termer."""
    sl = s.lower()
    has_obl = any(v in sl for v in GLOBAL_OBLIGATION_VERBS)
    has_units = bool(UNITS_REGEX.search(sl))
    has_terms = False
    if profile:
        has_units = has_units or bool(profile["units_re"].search(sl))
        has_terms = any(t in sl for t in profile.get("terms", []))
    return has_obl or has_units or has_terms

def _build_normative_block(sentences: list[str], start_idx: int, profile: dict | None = None, max_len_sent: int = 6) -> tuple[str, int]:
    """
    Bygger en 'normativ blokk' fra start_idx og nedover så lenge setningene er normative.
    Returnerer (tekst, siste_idx_inkl).
    """
    parts = [sentences[start_idx]]
    last = start_idx
    i = start_idx + 1
    while i < len(sentences) and (i - start_idx) < max_len_sent:
        nxt = sentences[i]
        if not _is_normative_sentence_profile(nxt, profile):
            break
        parts.append(nxt)
        last = i
        i += 1
    return " ".join(parts), last

def deduplicate_requirements(requirements: list, threshold: int = 95, scope: str = "per_file") -> list:
    """
    Slår sammen nesten-like krav til ett funn, med sammenslåtte referanser.
    - threshold: 0–100. 95 er strengt; 88–93 er ofte bra for like normative avsnitt.
    - scope: "per_file" (default) eller "global".
    """
    if not requirements:
        return []

    def key(req):
        return req['ref'].split(' / ')[0] if scope == "per_file" else "__GLOBAL__"

    groups = defaultdict(list)
    for r in requirements:
        groups[key(r)].append(r)

    out = []
    for _, reqs in groups.items():
        reqs_sorted = sorted(reqs, key=lambda x: (x.get('score', 0.0), len(x.get('text',''))), reverse=True)
        clusters = []   # [representant, [items...]]
        reps_norm = []  # normalisert anker/tekst for representant

        for r in reqs_sorted:
            tr = _norm_source(r)
            placed = False
            for ci, (_rep, items) in enumerate(clusters):
                sim = max(
                    token_set_ratio(tr, reps_norm[ci]),
                    fuzz_ratio(tr, reps_norm[ci]),
                    partial_ratio(tr, reps_norm[ci])
                )
                if sim >= threshold:
                    items.append(r)
                    placed = True
                    break
            if not placed:
                clusters.append([r, [r]])
                reps_norm.append(_norm_source(r))

        for rep, items in clusters:
            merged = dict(rep)
            merged['dup_count'] = len(items) - 1
            merged['ref'] = _merge_refs(items)

            # konsolider NS-treff
            if any('ns_treff' in it for it in items):
                ns_all, seen = [], set()
                for it in items:
                    ns_all.extend(it.get('ns_treff', []))
                dedup_ns = []
                for t in ns_all:
                    k = (t.get('standard'), t.get('side'), t.get('tekst'))
                    if k not in seen:
                        seen.add(k); dedup_ns.append(t)
                merged['ns_treff'] = dedup_ns

            if merged['dup_count'] > 0:
                merged['short_text'] = (merged.get('short_text') or "") + f" (konsolidert x{merged['dup_count']+1})"

            out.append(merged)

    return out

def create_reports_and_zip(requirements_list: list, temp_dir: Path, processing_errors: list, task_instance=None):
    """
    Gjenbrukbar funksjon som genererer alle rapporter (Excel, Word) basert på en
    liste med krav, og pakker dem i en ZIP-fil.
    """
    if task_instance:
        task_instance.update_state(state='PROGRESS', meta={'status': 'Genererer rapporter...', 'current': 85, 'total': 100})

    # Bygg datastrukturer som rapportfunksjonene forventer
    all_requirements = defaultdict(list)
    for req in requirements_list:
        file_name_part = req['ref'].split(' / ')[0]
        all_requirements[file_name_part].append(req)

    total_requirements_found = len(requirements_list)
    grouped_requirements = defaultdict(list)
    for req in requirements_list:
        grouped_requirements[req.get('gruppe', 'Uspesifisert')].append(req)

    keyword_counts, synonym_counts, semantic_search_counts = defaultdict(int), defaultdict(int), defaultdict(int)
    top_score_krav = None
    if requirements_list:
        top_score_krav = max(requirements_list, key=lambda x: x['score'])
        for req in requirements_list:
            if req['score'] >= 100.0: keyword_counts[req['keyword']] += 1
            elif req['score'] >= 95.0: synonym_counts[req['keyword']] += 1
            else: semantic_search_counts[req['keyword']] += 1
    
    total_files_scanned = len(all_requirements.keys())
    keywords_used = sorted(list(set(r['keyword'] for r in requirements_list if r['keyword'] != '(AI)')))
    
    try:
        excel_report = generate_total_excel_report_grouped(
            all_requirements, keywords_used, total_files_scanned, total_requirements_found,
            keyword_counts, synonym_counts, semantic_search_counts, grouped_requirements
        )
        delivery_report = generate_colored_delivery_report(grouped_requirements)
        graphical_report = generate_graphical_report_docx(
            total_requirements_found, keyword_counts, synonym_counts,
            semantic_search_counts, grouped_requirements, all_requirements
        )
        summary_report = generate_summary_report(
            all_requirements, keywords_used, total_files_scanned, total_requirements_found,
            keyword_counts, synonym_counts, top_score_krav, grouped_requirements,
            semantic_search_counts, []
        )

        excel_path = temp_dir / "Kravsporing_rapport.xlsx"
        word_delivery_path = temp_dir / "Leverandorleveranser_rapport.docx"
        summary_path = temp_dir / "Kravsporing_Oppsummering.docx"
        graphical_path = temp_dir / "Grafisk_Oversikt.docx"

        excel_path.write_bytes(excel_report.getvalue())
        word_delivery_path.write_bytes(delivery_report.getvalue())
        summary_path.write_bytes(summary_report.getvalue())
        graphical_path.write_bytes(graphical_report.getvalue())

        zip_final_path = temp_dir / "Kravsporing_Resultater.zip"
        with zipfile.ZipFile(zip_final_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            zipf.write(excel_path, arcname=excel_path.name)
            zipf.write(word_delivery_path, arcname=word_delivery_path.name)
            zipf.write(summary_path, arcname=summary_path.name)
            zipf.write(graphical_path, arcname=graphical_path.name)
            if processing_errors:
                error_report_path = temp_dir / "feilrapport.txt"
                error_report_path.write_text("\n".join(processing_errors), encoding='utf-8')
                zipf.write(error_report_path, arcname=error_report_path.name)
        
        if task_instance:
            task_instance.update_state(state='PROGRESS', meta={'status': 'Fullfører...', 'current': 100, 'total': 100})

    except Exception as e:
        logging.error(f"Feil under rapportgenerering i {temp_dir.name}: {e}", exc_info=True)
        if processing_errors is not None:
             processing_errors.append(f"Klarte ikke å generere rapporter: {e}")
        if task_instance:
            task_instance.update_state(state='FAILURE', meta={'status': f'Rapportfeil: {e}'})

@celery.task(bind=True, name="app.tasks.process_files_task")
def process_files_task(
    self,
    temp_dir_path: str,
    keywords,
    min_score,
    user_id: int,
    ns_standard_selection: str,
    mode: str,
    selected_groups,
    fokusomraade: str,
    ai_settings: dict,
):
    import logging, json
    from pathlib import Path

    log = logging.getLogger(__name__)

    temp_dir = Path(temp_dir_path)
    temp_id = temp_dir.name  # <-- NØKKEL: bruk denne overalt

    def _p(status: str, current: int):
        self.update_state(
            state="PROGRESS",
            meta={"status": status, "current": current, "total": 100, "temp_folder_id": temp_id}
        )

    # ---------- Init ----------
    _p("Starter...", 0)
    files_to_process = [f for f in temp_dir.iterdir() if f.is_file()]
    total_files = len(files_to_process)

    processing_errors: list[str] = []
    initial_requirements: list[dict] = []
    final_requirements: list[dict] = []

    if total_files == 0:
        log.error("Ingen filer å prosessere i %s", temp_dir)
        return {
            "status": "Ingen filer å behandle",
            "zip_folder": temp_id,
            "temp_folder_id": temp_id,  # <--
            "errors": ["Ingen filer å prosessere"],
            "preview": {"requirements": []},
        }

    # ---------- Hovedsløyfe ----------
    for idx, fpath in enumerate(files_to_process, start=1):
        _p(f"Behandler fil {idx}/{total_files}: {fpath.name}", 5 + int(60 * idx / max(1, total_files)))

        try:
            reqs, errs = _process_single_document(
                self_task=self,
                file_path=fpath,
                keywords=keywords or [],
                min_score=float(min_score),
                ns_standard_selection=ns_standard_selection,
                mode=mode,
                fokusomraade=fokusomraade or "",
                selected_groups=selected_groups or []
            )

            for r in reqs or []:
                if "fag" not in r:
                    best_fag, _ = classify_group_ai(r.get("text", "")) if r.get("text") else ("Uspesifisert", [])
                    r["fag"] = [best_fag or "Uspesifisert"]
                elif isinstance(r["fag"], str):
                    r["fag"] = [r["fag"]]
                if "status" not in r:
                    r["status"] = "Aktiv"
                st = r.get("short_text") or r.get("korttekst")
                if not st:
                    st = _generate_short_text(r.get("text", ""))
                r["short_text"] = st
                r["korttekst"] = st

            initial_requirements.extend(reqs or [])
            processing_errors.extend(errs or [])

        except Exception as e:
            msg = f"Feil ved behandling av {fpath.name}: {e}"
            log.error(msg, exc_info=True)
            processing_errors.append(msg)

    # ---------- Lagre rå funn ----------
    try:
        _p("Lagrer rå funn...", 70)
        initial_path = temp_dir / "initial_requirements.json"
        with open(initial_path, "w", encoding="utf-8") as f:
            json.dump(initial_requirements, f, ensure_ascii=False, indent=2)
    except Exception as e:
        processing_errors.append(f"Feil ved lagring av initial_requirements.json: {e}")

    # ---------- Etterbehandling ----------
    try:
        _p("Etterbehandler funn...", 80)
        filtered = [r for r in (initial_requirements or []) if float(r.get("score", 0.0)) >= float(min_score)]
        deduped = deduplicate_requirements(filtered, threshold=93, scope="per_file")
        final_requirements = _sort_requirements(deduped or [])  # KS-2 sortering
        if not isinstance(final_requirements, list):
            final_requirements = []
    except Exception as e:
        processing_errors.append(f"Feil i etterbehandling av krav: {e}")
        if not isinstance(final_requirements, list):
            final_requirements = []

    # ---------- Rapporter + ZIP ----------
    try:
        _p("Genererer rapporter og ZIP...", 90)
        create_reports_and_zip(final_requirements, temp_dir, processing_errors, self)
    except Exception as e:
        processing_errors.append(f"Feil ved generering av rapport/ZIP: {e}")

    # ---------- Ferdig ----------
    _p("Ferdigstiller...", 98)

    result_payload = {
        "status": "Rapport generert!",
        "zip_folder": temp_id,
        "temp_folder_id": temp_id,  # <--- VIKTIG
        "errors": processing_errors,
        "preview": {"requirements": final_requirements},
    }
    self.update_state(state="SUCCESS", meta={"status": "Fullført", "current": 100, "total": 100, "temp_folder_id": temp_id})
    return result_payload

# ===================================================================
# === STEG 3: Ny Celery-oppgave for å generere ZIP fra reviderte data ===
# ===================================================================
@celery.task(bind=True)
def generate_zip_from_review_task(self, user_id: int, temp_folder_id: str):
    """
    Laster reviderte krav (eller faller tilbake på originale),
    og genererer en ny ZIP-fil med rapporter.
    """
    self.update_state(state='PROGRESS', meta={'status': 'Starter generering...', 'current': 0, 'total': 100})

    temp_dir = TEMP_ROOT / temp_folder_id
    reviewed_path = temp_dir / "reviewed_requirements.json"
    initial_path = temp_dir / "initial_requirements.json"

    requirements_to_process: list[dict] = []
    errors: list[str] = []

    try:
        # --- Last datakilden ---
        if reviewed_path.exists():
            self.update_state(state='PROGRESS', meta={'status': 'Laster reviderte data...', 'current': 10, 'total': 100})
            with open(reviewed_path, 'r', encoding='utf-8') as f:
                requirements_to_process = json.load(f)
        elif initial_path.exists():
            self.update_state(state='PROGRESS', meta={'status': 'Laster originale data...', 'current': 10, 'total': 100})
            with open(initial_path, 'r', encoding='utf-8') as f:
                requirements_to_process = json.load(f)
        else:
            raise FileNotFoundError("Fant ingen krav-fil å generere rapport fra.")

        # --- Normaliser felter som UI/rapporter forventer ---
        for r in requirements_to_process or []:
            if "fag" not in r:
                best_fag, _ = classify_group_ai(r.get("text", "")) if r.get("text") else ("Uspesifisert", [])
                r["fag"] = [best_fag or "Uspesifisert"]
            elif isinstance(r["fag"], str):
                r["fag"] = [r["fag"]]
            if "status" not in r:
                r["status"] = "Aktiv"
            st = r.get("short_text") or r.get("korttekst")
            if not st:
                st = _generate_short_text(r.get("text", ""))
            r["short_text"] = st
            r["korttekst"] = st

        # --- KS-2: sorter krav og grupper før rapport/ZIP ---
        requirements_to_process = _sort_requirements(requirements_to_process or [])

        # Hvis du i denne funksjonen bygger/vedlikeholder gruppert struktur i et result-objekt,
        # bruk _sort_groups_order(keys) for å flytte 'Uspesifisert' til slutt. Eksempel:
        #
        # result = {"requirements": requirements_to_process, "by_fag": group_by_fag(requirements_to_process)}
        # if isinstance(result.get("by_fag"), dict):
        #     keys = list(result["by_fag"].keys())
        #     ordered = _sort_groups_order(keys)
        #     result["by_fag"] = {k: result["by_fag"][k] for k in ordered}

        # --- Generer rapport/ZIP ---
        self.update_state(state='PROGRESS', meta={'status': 'Genererer rapporter og ZIP...', 'current': 80, 'total': 100})
        create_reports_and_zip(requirements_to_process, temp_dir, errors, self)

    except Exception as e:
        logging.error(f"Feil under ZIP-generering fra review for {temp_folder_id}: {e}", exc_info=True)
        self.update_state(state='FAILURE', meta={'status': f'Kritisk feil: {e}'})
        raise e

    # NB: inkluderer zip_folder slik /status kan bygge riktig download_url
    return {
        'status': 'Fullført!',
        'zip_folder': temp_folder_id,
        'errors': errors
    }
 
# ===========================================================
# === Ny Celery-oppgave for å re-trene AI-modellene       ===
# =========================================================== 
@celery.task(bind=True, name="app.tasks.retrain_ai_task")
def retrain_ai_task(self, user_id: int, temp_folder_id: str):
    """
    Normaliserer datafiler til UTF-8, slår sammen review til korpus og kjører trening.
    Returnerer et resultat-objekt med merge-/train-status til frontend.
    """
    # 1) Normaliser alle .csv/.txt/.json i app/data
    DATA_DIR = Path("app") / "data"
    from app.routes.kravsporing import _normalize_data_dir, TEMP_ROOT, _merge_review_into_corpus, _run_training
    _normalize_data_dir(DATA_DIR)

    # 2) Lazy-import for å unngå sirkulær import
    from app.routes.kravsporing import TEMP_ROOT, _merge_review_into_corpus, _run_training

    # 3) Les reviewed/initial fra temp-mappe (hentet i web-UI)
    folder = (TEMP_ROOT / temp_folder_id)
    reviewed = folder / "reviewed_requirements.json"
    initial  = folder / "initial_requirements.json"
    src = reviewed if reviewed.exists() else initial

    merge_result = {"csv_added": 0, "csv_updated": 0, "neg_added": 0}
    try:
        if src.exists():
            import json
            reqs = json.loads(src.read_text(encoding="utf-8"))
            if isinstance(reqs, list):
                merge_result = _merge_review_into_corpus(reqs)
    except Exception as e:
        return {
            "status": "Fletting til korpus feilet",
            "errors": [str(e)],
            "temp_folder_id": temp_folder_id,
            "zip_folder": temp_folder_id,
        }

    # 4) Kjør trening
    try:
        train = _run_training()
        status = "AI-trening ferdig" if (train.get("fag_ok") and train.get("val_ok")) else "AI-trening ferdig med feil"
        errors = []
        if not train.get("fag_ok"): errors.append("Fag-trening feilet")
        if not train.get("val_ok"): errors.append("Validator-trening feilet")

        return {
            "status": status,
            "merge": merge_result,
            "train": train,
            "temp_folder_id": temp_folder_id,
            "zip_folder": temp_folder_id,
            "errors": errors,
        }
    except Exception as e:
        return {
            "status": "AI-trening feilet",
            "merge": merge_result,
            "errors": [str(e)],
            "temp_folder_id": temp_folder_id,
            "zip_folder": temp_folder_id,
        }

def get_ai_status():
    return {
        "nb_bert": {"ready": bool(NB_BERT_READY), "model": NB_BERT_NAME},
        "mnli": {"ready": bool(MNLI_READY), "model": MNLI_NAME},
    }