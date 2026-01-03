# app/tasks/reporting.py
# -*- coding: utf-8 -*-
"""
Dette modulet inneholder all logikk for å generere rapporter (Excel, Word)
og for å utheve nøkkelord i ulike filformater.
"""

# Setter Matplotlib til 'Agg' for å unngå GUI-problemer i bakgrunnstråder
import os
os.environ.setdefault("MPLBACKEND", "Agg")
try:
    import matplotlib
    matplotlib.use("Agg", force=True)
except Exception:
    pass

import io
import json
import re
import time
import zipfile
import logging
from pathlib import Path
from collections import defaultdict, OrderedDict

# Tredjepartsbiblioteker for rapportering
import fitz  # For PDF highlighting
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook, load_workbook
from openpyxl.styles import PatternFill, Alignment, Font, Border, Side
from openpyxl.utils import get_column_letter
from rapidfuzz.fuzz import ratio as fuzz_ratio
import matplotlib.pyplot as plt

# Importerer delt konfigurasjon
from app.config import GROUP_COLORS

# Definerer stier som er relevante for dette modulet
CURRENT_DIR = Path(__file__).resolve().parent.parent  # Peker til 'app'-mappen
TEMP_ROOT = CURRENT_DIR / "temp"
SYNONYM_PATH = CURRENT_DIR / "synonyms.json"
USPEC_LABELS = {"Uspesifisert", "uspesifisert", "Uspesifisert/ukjent", "Ukjent"}
# Sørg for at temp finnes når grafikk lagres
os.makedirs(TEMP_ROOT, exist_ok=True)


# === Hjelpefunksjoner for utheving ===

def _load_word_family_map():
    """Laster synonymer fra JSON for å utvide søketermer for utheving."""
    m = {}
    if SYNONYM_PATH.exists():
        try:
            with open(SYNONYM_PATH, "r", encoding="utf-8") as f:
                syns = json.load(f)
            for base, syn_list in syns.items():
                fam = [base.lower()] + [s.lower() for s in (syn_list or [])]
                for w in fam:
                    m[w] = fam
        except json.JSONDecodeError:
            pass
    return m


def highlight_docx(path, keywords, output):
    word_map = _load_word_family_map()
    doc = Document(path)
    keywords_lower = [k.lower() for k in (keywords or [])]
    all_terms = set(keywords_lower)
    for base in keywords_lower:
        if base in word_map:
            all_terms.update(word_map[base])
    terms = sorted([re.escape(t) for t in all_terms if t], key=len, reverse=True)
    if not terms:
        doc.save(output)
        return
    pattern = r'\b(' + '|'.join(terms) + r')\b'
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        rebuilt, last = [], 0
        for m in re.finditer(pattern, p.text, flags=re.IGNORECASE):
            if m.start() > last:
                rebuilt.append((p.text[last:m.start()], False))
            rebuilt.append((m.group(0), True))
            last = m.end()
        if last < len(p.text):
            rebuilt.append((p.text[last:], False))

        # Tøm eksisterende runs
        while p.runs:
            r = p.runs[0]
            r._r.getparent().remove(r._r)

        # Skriv inn igjen med highlight
        for txt, hi in rebuilt:
            if not txt:
                continue
            run = p.add_run(txt)
            if hi:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    doc.save(output)


def highlight_pdf(path, keywords, output):
    word_map = _load_word_family_map()
    doc = fitz.open(path)
    keywords_lower = [k.lower() for k in (keywords or [])]
    all_terms = set(keywords_lower)
    for base in keywords_lower:
        if base in word_map:
            all_terms.update(word_map[base])

    # Sett egnede flags hvis tilgjengelig i denne PyMuPDF-versjonen
    flags = 0
    if hasattr(fitz, "TEXT_IGNORECASE"):
        flags |= fitz.TEXT_IGNORECASE
    if hasattr(fitz, "TEXT_DEHYPHENATE"):
        flags |= fitz.TEXT_DEHYPHENATE

    for page in doc:
        for kw in all_terms:
            if not kw:
                continue
            try:
                rects = page.search_for(kw, flags=flags)
            except TypeError:
                # Eldre versjoner: fall tilbake uten flags
                rects = page.search_for(kw)
            for inst in rects:
                page.add_highlight_annot(inst).update()

    doc.save(output, garbage=4, deflate=True, clean=True)
    doc.close()


def highlight_excel(path, keywords, output):
    word_map = _load_word_family_map()
    wb = load_workbook(path)
    fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
    keywords_lower = [k.lower() for k in (keywords or [])]
    all_terms = set(keywords_lower)
    for base in keywords_lower:
        if base in word_map:
            all_terms.update(word_map[base])
    for s in wb.worksheets:
        for row in s.iter_rows():
            for c in row:
                if c.value is not None:
                    try:
                        if any(term in str(c.value).lower() for term in all_terms):
                            c.fill = fill
                    except Exception:
                        # robusthet på "rare" celler/typer
                        pass
    wb.save(output)


def highlight_text(path, keywords, output):
    word_map = _load_word_family_map()
    lines = Path(path).read_text(encoding='utf-8').splitlines()
    keywords_lower = [k.lower() for k in (keywords or [])]
    all_terms = set(keywords_lower)
    for base in keywords_lower:
        if base in word_map:
            all_terms.update(word_map[base])
    with open(output, "w", encoding="utf-8") as f:
        for line in lines:
            if any(term in line.lower() for term in all_terms):
                f.write(f"* {line.strip()}\n")
            else:
                f.write(f"{line.strip()}\n")


# === Rapport-genereringsfunksjoner ===

def generate_total_excel_report_grouped(
    all_requirements: dict,
    keywords: list,
    total_files_scanned: int,
    total_requirements_found: int,
    keyword_counts: defaultdict,
    synonym_counts: defaultdict,
    semantic_search_counts: defaultdict,
    grouped_requirements: defaultdict
) -> io.BytesIO:
    """
    Lager samlet Excel-rapport  ett ark per gruppe.
    """
    # Helper: trygt formattere NS-referanse fra krav["ns_pin"]
    def _fmt_ns_pin(krav: dict) -> str:
        pin = krav.get("ns_pin")
        if isinstance(pin, dict):
            std = (pin.get("standard") or "").strip()
            side = (pin.get("side") or "")
            if std or side:
                return f"{std} s.{side}" if side else std
        return ""

    # Helper: formatter "Forklaring" felt fra krav["explain"]
    def _fmt_explain(krav: dict) -> str:
        ex = krav.get("explain")
        if isinstance(ex, dict):
            try:
                kw  = float(ex.get("kw_sc", 0))
                sem = float(ex.get("sem_sc", 0))
                ai  = float(ex.get("ai_sc", 0))
                fok = float(ex.get("fokus_boost", 0))
            except Exception:
                kw = sem = ai = fok = 0.0
            return f"KW:{kw:.1f}% | SEM:{sem:.1f}% | AI:{ai:.1f}% | FOKUS:{fok:.1f}"
        return ""
       
    # ---------- Opprett workbook  standardoppsett ----------
    wb = Workbook()
    ws_all = wb.active
    ws_all.title = "Alle funn"
    headers = [
        "Søkeord","Korttekst","Funnet tekst","Kravtype","Treff %",
        "Dokument og side","Valgt løsning","Risiko","Kommentar",
        "Kravverb","# setninger","Lengde","NS-dokument(er)","NS-referanse","NS (valgt)","NS-kvalitet","Forklaring","Topp fag","Gruppe"
    ]
    ws_all.append(headers)
    bold_font = Font(bold=True)
    gray_fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )

    # ---------- Normaliser innkommende krav til en flat liste ----------
    # Tillat at data kan komme som:
    # - grouped_requirements (defaultdict[str, list[dict]])
    # - all_requirements={"requirements":[...]} eller {"preview":{"requirements":[...]}}
    # - eller direkte en liste under all_requirements (fallback)
    def _flatten_grouped(g):
        flat = []
        for grp, items in (g or {}).items():
            for it in items or []:
                if "gruppe" not in it or not it.get("gruppe"):
                    # fallback til første fag som gruppenavn
                    fag = it.get("fag") or []
                    it["gruppe"] = (fag[0] if isinstance(fag, list) and fag else "Uspesifisert")
                flat.append(it)
        return flat

    if grouped_requirements and isinstance(grouped_requirements, dict):
        requirements = _flatten_grouped(grouped_requirements)
    else:
        requirements = []
        if isinstance(all_requirements, dict):
            if "requirements" in all_requirements and isinstance(all_requirements["requirements"], list):
                requirements = all_requirements["requirements"]
            elif "preview" in all_requirements and isinstance(all_requirements["preview"], dict):
                reqs = all_requirements["preview"].get("requirements")
                if isinstance(reqs, list):
                    requirements = reqs
                # OBS: 'preview' brukes nå kun til eget ark "Usikre funn"; hoved-arket bygger på 'requirements'.
        if not requirements and isinstance(all_requirements, list):
            requirements = all_requirements

        # Sett gruppe ut fra fag dersom gruppe mangler
        for it in requirements or []:
            if "gruppe" not in it or not it.get("gruppe"):
                fag = it.get("fag") or []
                it["gruppe"] = (fag[0] if isinstance(fag, list) and fag else "Uspesifisert")

    # ---------- Fyll "Alle funn"-arket ----------
    for krav in (requirements or []):
        nlp_v = krav.get("nlp_vurdering", {}) or {}
        kravverb = "Ja" if nlp_v.get("inneholder_kravsverb") else "Nei"
        ant_setn = nlp_v.get("antall_setninger", "")
        ant_tegn = nlp_v.get("lengde", "")

        ns_treff = krav.get("ns_treff") or []
        std_dok = ", ".join(sorted({t.get("standard") for t in ns_treff if t.get("standard")})) if ns_treff else ""

        utdrag = []
        for t in ns_treff:
            tekst = (t.get("tekst") or "")
            preview = tekst[:100].replace("\n", " ").strip() + ("..." if len(tekst) > 100 else "")
            side = t.get("side", "")
            sc = t.get("score")
            try:
                sc_txt = f"{float(sc):.1f}%"
            except Exception:
                sc_txt = ""
            std = t.get("standard") or ""
            score_suffix = f" [Score: {sc_txt}]" if sc_txt else ""
            utdrag.append(f"{std} (s.{side}): {preview}{score_suffix}")

        ns_ref = "\n".join(utdrag) if utdrag else "Ingen relevante treff."

        ws_all.append([
            krav.get("keyword", ""),
            (krav.get("short_text") or krav.get("korttekst") or ""),
            krav.get("text", ""),
            krav.get("kravtype", ""),
            float(krav.get("score", 0.0)) / 100.0,  # lagres som 0–1 for Excel-prosentformat
            krav.get("ref", ""),
            "→ Beskriv valgt løsning her",
            "",
            "",
            kravverb,
            ant_setn,
            ant_tegn,
            std_dok,
            ns_ref,
            # NS (valgt)  NS-kvalitet
            _fmt_ns_pin(krav),
            (krav.get("ns_quality") or ""),
            # Forklaring  topp fag
            _fmt_explain(krav),
            (
                ", ".join(
                    f"{f.get('label','?')} ({f.get('score',0):.1f}%)"
                    for f in krav.get("top_fag", [])
                ) if isinstance(krav.get("top_fag"), list) else ""
            ),
            krav.get("gruppe", "Uspesifisert"),
        ])

    # ---------- Formatering for "Alle funn" ----------
    for r in ws_all.iter_rows(min_row=2, max_row=ws_all.max_row, min_col=1, max_col=len(headers)):
        for c in r:
            c.alignment = Alignment(wrap_text=True, vertical="top")
            c.border = Border(
                left=Side(style="thin"), right=Side(style="thin"),
                top=Side(style="thin"), bottom=Side(style="thin")
            )
    # Header: fet  grå, kolonnebredder
    for i, h in enumerate(headers, start=1):
        cell = ws_all.cell(row=1, column=i)
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
        width = (
            50 if h in ["Funnet tekst", "Valgt løsning", "Kommentar", "NS-referanse"]
            else 40 if h in ["Korttekst","NS (valgt)"]
            else 30 if h == "Dokument og side"
            else 22
        )
        ws_all.column_dimensions[get_column_letter(i)].width = width

    # Prosentformat på "Treff %" (kolonne 5)
    for row in range(2, ws_all.max_row + 1):
        ws_all.cell(row=row, column=5).number_format = "0.0%"

    ws_all.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{ws_all.max_row}"

    # ---------- Ark for "Usikre funn" (preview) ----------
    preview_items = []
    if isinstance(all_requirements, dict) and isinstance(all_requirements.get("preview"), dict):
        pr = all_requirements["preview"].get("requirements")
        if isinstance(pr, list):
            preview_items = pr

    if preview_items:
        ws_prev = wb.create_sheet("Usikre funn")
        ws_prev.append(headers)
        for krav in preview_items:
            nlp_v = krav.get("nlp_vurdering", {}) or {}
            kravverb = "Ja" if nlp_v.get("inneholder_kravsverb") else "Nei"
            ant_setn = nlp_v.get("antall_setninger", "")
            ant_tegn = nlp_v.get("lengde", "")
            ns_treff = krav.get("ns_treff") or []
            std_dok = ", ".join(sorted({t.get("standard") for t in ns_treff if t.get("standard")})) if ns_treff else ""

            utdrag = []
            for t in ns_treff:
                tekst = (t.get("tekst") or "")
                preview = tekst[:100].replace("\n", " ").strip() + ("..." if len(tekst) > 100 else "")
                side = t.get("side", "")
                sc = t.get("score")
                try:
                    sc_txt = f"{float(sc):.1f}%"
                except Exception:
                    sc_txt = ""
                std = t.get("standard") or ""
                score_suffix = f" [Score: {sc_txt}]" if sc_txt else ""
                utdrag.append(f"{std} (s.{side}): {preview}{score_suffix}")

            ns_ref = "\n".join(utdrag) if utdrag else "Ingen relevante treff."

            ws_prev.append([
                krav.get("keyword", ""),
                (krav.get("short_text") or krav.get("korttekst") or ""),
                krav.get("text", ""),
                krav.get("kravtype", ""),
                float(krav.get("score", 0.0)) / 100.0,
                krav.get("ref", ""),
                "→ Beskriv valgt løsning her",
                "",
                "",
                kravverb,
                ant_setn,
                ant_tegn,
                std_dok,
                ns_ref,
                # NS (valgt)  NS-kvalitet
                _fmt_ns_pin(krav),
                (krav.get("ns_quality") or ""),
                _fmt_explain(krav),
                (
                    ", ".join(
                        f"{f.get('label','?')} ({f.get('score',0):.1f}%)"
                        for f in krav.get("top_fag", [])
                    ) if isinstance(krav.get("top_fag"), list) else ""
                ),
                krav.get("gruppe", "Uspesifisert"),
            ])

        # Formatering "Usikre funn"
        for r in ws_prev.iter_rows(min_row=1, max_row=ws_prev.max_row, min_col=1, max_col=len(headers)):
            for c in r:
                c.alignment = Alignment(wrap_text=True, vertical="top")
                c.border = thin_border
        for i, h in enumerate(headers, start=1):
            cell = ws_prev.cell(row=1, column=i)
            cell.font = bold_font
            cell.fill = gray_fill
            width = 50 if h in ["Funnet tekst", "Valgt løsning", "Kommentar", "NS-referanse"] \
                    else 40 if h in ["Korttekst","NS (valgt)"] \
                    else 30 if h == "Dokument og side" \
                    else 22
            ws_prev.column_dimensions[get_column_letter(i)].width = width
        for row in range(2, ws_prev.max_row + 1):
            ws_prev.cell(row=row, column=5).number_format = "0.0%"
        ws_prev.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{ws_prev.max_row}"

    # ---------- Ett ark per gruppe ----------
    def _group_key(it: dict) -> str:
        g = it.get("gruppe")
        if not g:
            fag = it.get("fag") or []
            return fag[0] if isinstance(fag, list) and fag else "Uspesifisert"
        return g

    grupper = sorted({ _group_key(it) for it in (requirements or []) }, key=lambda k: (k == "Uspesifisert", k))

    for group_name in grupper:
        title = (group_name or "Uspesifisert")[:31] or "Uspesifisert"
        ws_g = wb.create_sheet(title)
        ws_g.append(headers)

        group_rows = [it for it in (requirements or []) if _group_key(it) == group_name]

        for krav in group_rows:
            nlp_v = krav.get("nlp_vurdering", {}) or {}
            kravverb = "Ja" if nlp_v.get("inneholder_kravsverb") else "Nei"
            ant_setn = nlp_v.get("antall_setninger", "")
            ant_tegn = nlp_v.get("lengde", "")

            ns_treff = krav.get("ns_treff") or []
            std_dok = ", ".join(sorted({t.get("standard") for t in ns_treff if t.get("standard")})) if ns_treff else ""

            utdrag = []
            for t in ns_treff:
                tekst = (t.get("tekst") or "")
                preview = tekst[:100].replace("\n", " ").strip() + ("..." if len(tekst) > 100 else "")
                side = t.get("side", "")
                sc = t.get("score")
                try:
                    sc_txt = f"{float(sc):.1f}%"
                except Exception:
                    sc_txt = ""
                std = t.get("standard") or ""
                score_suffix = f" [Score: {sc_txt}]" if sc_txt else ""
                utdrag.append(f"{std} (s.{side}): {preview}{score_suffix}")

            ns_ref = "\n".join(utdrag) if utdrag else "Ingen relevante treff."

            ws_g.append([
                krav.get("keyword", ""),
                (krav.get("short_text") or krav.get("korttekst") or ""),
                krav.get("text", ""),
                krav.get("kravtype", ""),
                float(krav.get("score", 0.0)) / 100.0,
                krav.get("ref", ""),
                "→ Beskriv valgt løsning her",
                "",
                "",
                kravverb,
                ant_setn,
                ant_tegn,
                std_dok,
                ns_ref,
                # NS (valgt)  NS-kvalitet
                _fmt_ns_pin(krav),
                (krav.get("ns_quality") or ""),
                _fmt_explain(krav),
                (
                    ", ".join(
                        f"{f.get('label','?')} ({f.get('score',0):.1f}%)"
                        for f in krav.get("top_fag", [])
                    ) if isinstance(krav.get("top_fag"), list) else ""
                ),
                group_name or "Uspesifisert",
            ])

        # Formatering for gruppearket
        # 1) Header
        for i, h in enumerate(headers, start=1):
            cell = ws_g.cell(row=1, column=i)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color="DDDDDD", end_color="DDDDDD", fill_type="solid")
            width = (
                50 if h in ["Funnet tekst", "Valgt løsning", "Kommentar", "NS-referanse"]
                else 40 if h in ["Korttekst","NS (valgt)"]
                else 30 if h == "Dokument og side"
                else 22
            )
            ws_g.column_dimensions[get_column_letter(i)].width = width
        # 2) Celler
        for r in ws_g.iter_rows(min_row=2, max_row=ws_g.max_row, min_col=1, max_col=len(headers)):
            for c in r:
                c.alignment = Alignment(wrap_text=True, vertical="top")
                c.border = Border(
                    left=Side(style="thin"), right=Side(style="thin"),
                    top=Side(style="thin"), bottom=Side(style="thin")
                )

        # Prosentformat i grupper (kolonne 5)
        for row in range(2, ws_g.max_row + 1):
            ws_g.cell(row=row, column=5).number_format = "0.0%"

        ws_g.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{ws_g.max_row}"

    # ---------- Skriv ut workbook ----------
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def generate_colored_delivery_report(grouped_requirements: defaultdict) -> io.BytesIO:
    doc = Document()
    doc.add_heading("Leverandørleveranser", 0)
    doc.add_paragraph(
        "Denne rapporten grupperer funne krav basert på potensielt ansvarlig fagområde/leverandør. "
        "Verifiser grupperingen før bruk."
    )

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text, hdr[5].text = \
        "Gruppe", "Krav", "Type", "Referanse", "Forklaring", "Topp fag"

    # Egen shading-node per headercelle
    for c in hdr:
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), "C0C0C0")
        c._tc.get_or_add_tcPr().append(shd)

    def filtrer_unike(kravliste: list, terskel=95):
        unike = []
        for item in kravliste:
            if not any(fuzz_ratio((item.get('text') or '').lower(), (ex.get('text') or '').lower()) >= terskel for ex in unike):
                unike.append(item)
        return unike

    for grp in sorted(grouped_requirements.keys(), key=lambda g: (g == "Uspesifisert", g)):
        for krav in sorted(filtrer_unike(grouped_requirements[grp]), key=lambda x: x.get('score', 0), reverse=True):
            row = table.add_row().cells
            # Basis
            row[0].text = grp
            row[1].text = krav.get("text", "")
            row[2].text = krav.get("kravtype", "")
            row[3].text = krav.get("ref", "")
            # Forklaring
            ex = krav.get("explain") or {}
            if isinstance(ex, dict) and ex:
                row[4].text = (
                    f"KW:{float(ex.get('kw_sc',0)):.1f}% | "
                    f"SEM:{float(ex.get('sem_sc',0)):.1f}% | "
                    f"AI:{float(ex.get('ai_sc',0)):.1f}% | "
                    f"FOKUS:{float(ex.get('fokus_boost',0)):.1f}"
                )
            else:
                row[4].text = ""
            # Topp fag
            tf = krav.get("top_fag")
            if isinstance(tf, list) and tf:
                row[5].text = ", ".join(
                    f"{(f.get('label') or '?')} ({float(f.get('score',0)):.1f}%)" for f in tf
                )
            else:
                row[5].text = ""
            if grp in GROUP_COLORS:
                # Viktig: lag ny shading-node for hver celle (ikke gjenbruk samme node)
                for cell_in_row in row:
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), GROUP_COLORS[grp])
                    cell_in_row._tc.get_or_add_tcPr().append(shd)

    out = io.BytesIO()
    # Juster typografi for "Forklaring" og "Topp fag" (kolonner 5 og 6) til liten/grå
    for r_i, row in enumerate(table.rows):
        if r_i == 0:
            continue  # hopp header
        for c_i in (4, 5):  # 0-indeksert: 5. og 6. kolonne
            for p in row.cells[c_i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(120, 120, 120)
    doc.save(out)
    out.seek(0)
    return out


def generate_graphical_report_docx(
    total_requirements_found: int, keyword_counts: defaultdict,
    synonym_counts: defaultdict, semantic_search_counts: defaultdict,
    grouped_requirements: defaultdict, all_requirements: dict
) -> io.BytesIO:
    tmp_paths = []  # Why: rydde opp midlertidige bilder etter lagring
    plt.ioff()
    doc = Document()
    doc.add_heading('Grafisk Kravsporingsoversikt', 0)
    doc.add_paragraph(f"Rapport generert: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Totalt antall unike krav identifisert: {total_requirements_found}")
    doc.add_paragraph("")

    ts = int(time.time())

    # --- Fordeling av trefftyper (pie) ---
    try:
        labels_pie = ['Eksakte Nøkkelord', 'Eksakte Synonymer', 'Maskinlæresøk (AI-søk)']
        sizes_pie = [sum(keyword_counts.values()), sum(synonym_counts.values()), sum(semantic_search_counts.values())]
        if sum(sizes_pie) > 0:
            fig1, ax1 = plt.subplots(figsize=(8, 8))
            ax1.pie(sizes_pie, labels=labels_pie, autopct='%1.1f%%', startangle=90)
            ax1.axis('equal')
            ax1.set_title('Fordeling av Trefftyper')
            pie_path = TEMP_ROOT / f"trefftyper_pie_{ts}.png"
            tmp_paths.append(pie_path)
            plt.savefig(pie_path, bbox_inches='tight')
            plt.close(fig1)
            doc.add_heading('Fordeling av Trefftyper', level=1)
            doc.add_picture(str(pie_path), width=Inches(6))
            doc.add_paragraph("")
    except Exception as e:
        logging.warning(f"Kunne ikke lage pie-chart: {e}", exc_info=True)

    # --- Fordeling av kravtyper (bar) ---
    try:
        kravtype_counts = defaultdict(int)
        for krav_list in all_requirements.values():
            for item in krav_list:
                t = item.get('kravtype', 'krav')
                kravtype_counts[t] += 1
        if kravtype_counts:
            labels_bar, counts_bar = list(kravtype_counts.keys()), list(kravtype_counts.values())
            fig2, ax2 = plt.subplots(figsize=(10, 6))
            ax2.bar(labels_bar, counts_bar)
            ax2.set_ylabel('Antall krav')
            ax2.set_title('Fordeling av Kravtyper')
            ax2.tick_params(axis='x', rotation=45)
            bar_path = TEMP_ROOT / f"kravtyper_bar_{ts}.png"
            tmp_paths.append(bar_path)
            plt.tight_layout()
            plt.savefig(bar_path, bbox_inches='tight')
            plt.close(fig2)
            doc.add_heading('Fordeling av Kravtyper', level=1)
            doc.add_picture(str(bar_path), width=Inches(6))
            doc.add_paragraph("")
    except Exception as e:
        logging.warning(f"Kunne ikke lage bar-chart (kravtyper): {e}", exc_info=True)

    # --- Fordeling per gruppe (bar) ---
    try:
        if grouped_requirements:
            order = sorted(grouped_requirements.keys(), key=lambda g: (g == "Uspesifisert", g))
            labels_group = [g.capitalize() for g in order]
            counts_group = [len(grouped_requirements[g]) for g in order]
            fig3, ax3 = plt.subplots(figsize=(10, 6))
            ax3.bar(labels_group, counts_group)
            ax3.set_ylabel('Antall krav')
            ax3.set_title('Fordeling av Fagområder / Leverandørgrupper')
            ax3.tick_params(axis='x', rotation=45)
            grp_path = TEMP_ROOT / f"grupper_bar_{ts}.png"
            tmp_paths.append(grp_path)
            plt.tight_layout()
            plt.savefig(grp_path, bbox_inches='tight')
            plt.close(fig3)
            doc.add_heading('Fordeling av Fagområder / Leverandørgrupper', level=1)
            doc.add_picture(str(grp_path), width=Inches(6))
            doc.add_paragraph("")
    except Exception as e:
        logging.warning(f"Kunne ikke lage bar-chart (grupper): {e}", exc_info=True)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    try:
        for p in tmp_paths:
            try:
                os.unlink(p)
            except Exception:
                pass
    except Exception:
        pass
    return out


def generate_summary_report(
    all_requirements: dict, keywords: list, total_files_scanned: int,
    total_requirements_found: int, keyword_counts: defaultdict,
    synonym_counts: defaultdict, top_score_krav: dict,
    grouped_requirements: defaultdict, semantic_search_counts: defaultdict,
    suggested_keywords: list = None
) -> io.BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    doc.add_heading('Oppsummering av utført kravsporing', 0)
    doc.add_paragraph(f"Generert: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Nøkkelord brukt i søket: {', '.join(keywords or [])}" if keywords else "Analysemodus: Bare AI-modell (ingen nøkkelord).")
    doc.add_paragraph("")

    doc.add_heading('Søke- og funnstatistikk', 1)
    doc.add_paragraph(f"Totalt antall dokumenter skannet: {total_files_scanned}")
    doc.add_paragraph(f"Totalt antall unike krav identifisert: {total_requirements_found}")
    doc.add_paragraph("")

    doc.add_heading('Funn per kategori', 2)
    if keyword_counts:
        doc.add_paragraph("Eksakte nøkkelord:")
        for kw, c in sorted(keyword_counts.items(), key=lambda x: x[1], reverse=True):
            doc.add_paragraph(f"- '{kw}': {c} treff")
    else:
        doc.add_paragraph("Ingen eksakte nøkkelordstreff.")
    doc.add_paragraph("")

    if synonym_counts:
        doc.add_paragraph("Eksakte synonymer:")
        for kw, c in sorted(synonym_counts.items(), key=lambda x: x[1], reverse=True):
            doc.add_paragraph(f"- '{kw}': {c} treff")
    else:
        doc.add_paragraph("Ingen synonymtreff.")
    doc.add_paragraph("")

    if semantic_search_counts:
        doc.add_paragraph("Maskinlæresøk (AI):")
        for kw, c in sorted(semantic_search_counts.items(), key=lambda x: x[1], reverse=True):
            doc.add_paragraph(f"- '{kw}': {c} treff")
    else:
        doc.add_paragraph("Ingen AI-treff.")
    doc.add_paragraph("")

    if top_score_krav:
        doc.add_heading('Høyest skårende funn', 2)
        doc.add_paragraph(f"«{top_score_krav.get('text','')}»")
        doc.add_paragraph(f"Referanse: {top_score_krav.get('ref','')}")
        doc.add_paragraph(f"Søkeord: {top_score_krav.get('keyword','')} – Skår: {float(top_score_krav.get('score',0)):.1f}")
        doc.add_paragraph("")

    if grouped_requirements:
        doc.add_heading('Fordeling per fag/gruppe', 2)
        for g in sorted(grouped_requirements.keys(), key=lambda k: (k == "Uspesifisert", k)):
            doc.add_paragraph(f"- {g}: {len(grouped_requirements[g])} krav")

    if suggested_keywords:
        doc.add_paragraph("")
        doc.add_heading('Foreslåtte nye nøkkelord', 2)
        for s in suggested_keywords:
            doc.add_paragraph(f"- {s}")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out


# === Hovedfunksjon for å lage alle rapporter og ZIP-fil ===
def _group_requirements_by_fag(requirements: list[dict]) -> "OrderedDict[str, list[dict]]":
    tmp = defaultdict(list)
    for r in requirements or []:
        # Normaliser til streng: foretrekk eksplisitt 'gruppe', ellers første element i 'fag' (om liste)
        g = r.get("gruppe")
        if not g:
            fag = r.get("fag")
            if isinstance(fag, list) and fag:
                g = fag[0]
            elif isinstance(fag, str) and fag.strip():
                g = fag.strip()
            else:
                g = "Uspesifisert"
        tmp[g].append(r)
    # Sorter alfabetisk, men med Uspesifisert sist
    def _key(k: str) -> tuple[int, str]:
        return (1 if k in USPEC_LABELS else 0, k.lower())
    ordered = OrderedDict()
    for k in sorted(tmp.keys(), key=_key):
        # Sorter internt: søkeord A–Å, deretter treff% synkende
        ordered[k] = sorted(
            tmp[k],
            key=lambda x: ((x.get("keyword") or "").lower(), -float(x.get("score") or 0.0))
        )
    return ordered

def _autosize(ws):
    for col in ws.columns:
        maxlen = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            try:
                maxlen = max(maxlen, len(str(cell.value or "")))
            except Exception:
                pass
        ws.column_dimensions[col_letter].width = min(maxlen + 2, 60)

def _write_excel(path: Path, rows: list[dict]):
    PERCENT_COL_IDX = 4  # 1-basert indeks for "Treff %"
    # Hjelper: konverter lister/dict til lesbar tekst for Excel
    def _as_text(v):
        if v is None:
            return ""
        if isinstance(v, list):
            return ", ".join(str(x) for x in v if x is not None)
        if isinstance(v, dict):
            # kompakt JSON uten ASCII-escaping
            return json.dumps(v, ensure_ascii=False, separators=(",", ":"))
        return str(v)
    wb = Workbook()
    ws = wb.active
    ws.title = "Krav"
    headers = [
        "søkeord", "funnet tekst", "kravtype", "Treff %", "Dokument og side",
        "Valgt løsning", "Risiko", "Kommentar", "Fag", "NS-referanser", "NLP-vurdering"
    ]
    ws.append(headers)
    # Grå bakgrunn i kolonne A (søkeord)
    grey = PatternFill(start_color="00DDDDDD", end_color="00DDDDDD", fill_type="solid")
    for r in rows:
        # Normaliser problemfelter til strenger
        fag_val = r.get("gruppe") or r.get("fag") or "Uspesifisert"
        if isinstance(fag_val, list):
            fag_val = ", ".join(str(x) for x in fag_val if x)
        ns_refs_txt = _as_text(r.get("ns_refs", ""))
        nlp_eval_txt = _as_text(r.get("nlp_eval", ""))
        ws.append([
            r.get("keyword", ""),
            r.get("match_text", "")[:5000],
            r.get("kravtype", ""),
            round(float(r.get("score", 0.0)) / 100.0, 4),
            r.get("ref", ""),
            "** Velg løsning **" if r.get("keyword") else "",
            r.get("risiko", ""),
            r.get("kommentar", ""),
            fag_val,
            ns_refs_txt,
            nlp_eval_txt,
        ])
    # Formatering
    for cell in ws["A"][1:]:
        cell.fill = grey
    ws.auto_filter.ref = ws.dimensions
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    _autosize(ws)
    # Prosentformat på "Treff %"
    for row in range(2, ws.max_row + 1):
        ws.cell(row=row, column=PERCENT_COL_IDX).number_format = "0.0%"
    wb.save(path)

def export_per_fag_excel_zip(requirements: list[dict], out_dir: Path) -> Path:
    """
    Lager per-fag Excel-filer i underkatalog /per_fag og returnerer sti til ZIP.
    Endrer ingenting i eksisterende hovedrapportflyt – kan kalles i tillegg.
    """
    out_dir.mkdir(parents=True, exist_ok=True)
    per_fag_dir = out_dir / "per_fag"
    per_fag_dir.mkdir(exist_ok=True)

    grouped = _group_requirements_by_fag(requirements)
    produced = []
    for fag, rows in grouped.items():
        safe = re.sub(r"[^0-9A-Za-zæøåÆØÅ._ -]+", "_", fag).strip().strip("_")
        path = per_fag_dir / f"krav_{safe or 'Uspesifisert'}.xlsx"
        _write_excel(path, rows)
        produced.append(path)

    # ZIP
    import zipfile
    zip_path = out_dir / "per_fag_rapporter.zip"
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for p in produced:
            zf.write(p, arcname=str(p.relative_to(out_dir)))
    return zip_path
    
def create_reports_and_zip(
    requirements_list: list,
    temp_dir: Path,
    processing_errors: list,
    task_instance=None
):
    """
    Genererer alle rapporter og én samlet ZIP i temp_dir.
    Kaster Exception ved kritiske feil, slik at Celery markerer FAILURE.
    """
    log = logging.getLogger(__name__)

    def _progress(status: str, current: int):
        if task_instance:
            try:
                task_instance.update_state(
                    state="PROGRESS",
                    meta={"status": status, "current": int(current), "total": 100},
                )
            except Exception:
                log.warning("update_state feilet (ignorerer).", exc_info=True)

    _progress("Forbereder rapportdata…", 85)

    # --- Samle metadata ---
    # per-kilde (filnavn-delen av ref)
    all_requirements = defaultdict(list)
    for req in requirements_list or []:
        file_name_part = (req.get("ref") or "").split(" / ")[0]
        all_requirements[file_name_part].append(req)

    total_requirements_found = len(requirements_list or [])
    total_files_scanned = len(all_requirements.keys())

    # Grupper per 'gruppe' (fallback: første fag, ellers 'Uspesifisert')
    grouped_requirements = defaultdict(list)
    for req in requirements_list or []:
        grp = req.get("gruppe")
        if not grp:
            fag = req.get("fag") or []
            grp = fag[0] if isinstance(fag, list) and fag else "Uspesifisert"
            req["gruppe"] = grp
        grouped_requirements[grp].append(req)

    # Enkle «score buckets» (behold logikken din)
    keyword_counts = defaultdict(int)
    synonym_counts = defaultdict(int)
    semantic_search_counts = defaultdict(int)

    def _sc(it):
        try:
            return float(it.get("score", 0.0))
        except Exception:
            return 0.0
    top_score_krav = max(requirements_list, key=_sc) if requirements_list else None
    for req in requirements_list or []:
        sc = float(req.get("score", 0.0))
        kw = req.get("keyword", "")
        if sc >= 100.0:
            keyword_counts[kw] += 1
        elif sc >= 95.0:
            synonym_counts[kw] += 1
        else:
            semantic_search_counts[kw] += 1

    keywords_used = sorted(
        {r.get("keyword") for r in (requirements_list or []) if r.get("keyword") and r.get("keyword") != "(AI)"}
    )

    # --- Generer delrapporter (kaster ved feil) ---
    try:
        _progress("Genererer Excel-rapport…", 90)
        # Merk: generate_total_excel_report_grouped forventer en struktur som vi normaliserte i tidligere revisjon.
        excel_report = generate_total_excel_report_grouped(
            {"requirements": requirements_list},
            keywords_used,
            total_files_scanned,
            total_requirements_found,
            keyword_counts,
            synonym_counts,
            semantic_search_counts,
            grouped_requirements,
        )

        _progress("Genererer leveranse-rapport…", 92)
        delivery_report = generate_colored_delivery_report(grouped_requirements)

        _progress("Genererer grafisk rapport…", 94)
        graphical_report = generate_graphical_report_docx(
            total_requirements_found,
            keyword_counts,
            synonym_counts,
            semantic_search_counts,
            grouped_requirements,
            all_requirements,
        )

        _progress("Genererer oppsummering…", 96)
        summary_report = generate_summary_report(
            all_requirements,
            keywords_used,
            total_files_scanned,
            total_requirements_found,
            keyword_counts,
            synonym_counts,
            top_score_krav,
            grouped_requirements,
            semantic_search_counts,
            [],  # evt. reserved arg for fremtidige avvik/varsler
        )

    except Exception as e:
        msg = f"Feil under rapportgenerering: {e}"
        log.error(msg, exc_info=True)
        if processing_errors is not None:
            processing_errors.append(msg)
        # Kritisk: ikke fortsett til ZIP uten rapporter
        raise

    # --- Skriv filer til disk ---
    try:
        _progress("Lagrer rapportfiler…", 98)
        (temp_dir / "Kravsporing_rapport.xlsx").write_bytes(excel_report.getvalue())
        (temp_dir / "Leverandorleveranser_rapport.docx").write_bytes(delivery_report.getvalue())
        (temp_dir / "Kravsporing_Oppsummering.docx").write_bytes(summary_report.getvalue())
        (temp_dir / "Grafisk_Oversikt.docx").write_bytes(graphical_report.getvalue())
        # Generer og lagre per-fag-rapporter (valgfri ekstra leveranse)
        try:
            per_fag_zip = export_per_fag_excel_zip(requirements_list, temp_dir)
        except Exception as e:
            per_fag_zip = None
            log.warning(f"Per-fag-eksport feilet: {e}", exc_info=True)
    except Exception as e:
        msg = f"Feil ved skriving av rapportfiler: {e}"
        log.error(msg, exc_info=True)
        if processing_errors is not None:
            processing_errors.append(msg)
        raise

    # --- Bygg ZIP ---
    try:
        _progress("Pakker ZIP…", 99)
        zip_final_path = temp_dir / "Kravsporing_Resultater.zip"
        with zipfile.ZipFile(zip_final_path, "w", zipfile.ZIP_DEFLATED) as zipf:
            zipf.write(temp_dir / "Kravsporing_rapport.xlsx", arcname="Kravsporing_rapport.xlsx")
            zipf.write(temp_dir / "Leverandorleveranser_rapport.docx", arcname="Leverandorleveranser_rapport.docx")
            zipf.write(temp_dir / "Kravsporing_Oppsummering.docx", arcname="Kravsporing_Oppsummering.docx")
            zipf.write(temp_dir / "Grafisk_Oversikt.docx", arcname="Grafisk_Oversikt.docx")
            # Ta med per-fag ZIP hvis produsert
            if (temp_dir / "per_fag_rapporter.zip").exists():
                zipf.write(temp_dir / "per_fag_rapporter.zip", arcname="per_fag_rapporter.zip")
            if processing_errors:
                error_report_path = temp_dir / "feilrapport.txt"
                error_report_path.write_text("\n".join(processing_errors), encoding="utf-8")
                zipf.write(error_report_path, arcname=error_report_path.name)

        _progress("Fullfører…", 100)
        return zip_final_path  # Nyttig for e2e-tester og kallere
    except Exception as e:
        msg = f"Feil under ZIP-generering: {e}"
        log.error(msg, exc_info=True)
        if processing_errors is not None:
            processing_errors.append(msg)
        raise
