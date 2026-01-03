# -*- coding: utf-8 -*-
"""
Dette modulet inneholder all logikk for å hente ut ren tekst fra ulike
fildokumenter (PDF, DOCX, DOC, XLSX, TXT, CSV, MSG).
"""
from __future__ import annotations

import logging
import shutil
import subprocess
import io
import secrets
import re
from pathlib import Path

# Tredjepartsbiblioteker for fil-parsing
import fitz  # PyMuPDF for PDF
import extract_msg  # For .msg (Outlook-eposter)
from docx import Document  # For .docx
from openpyxl import load_workbook  # For .xlsx

from .core import extract_requirements, clean_text

log = logging.getLogger(__name__)

# Definerer TEMP_ROOT her også, slik at modulen er selvstendig
# Antar at denne filen er i app/tasks/, så vi går to nivåer opp for å finne app-roten
TEMP_ROOT = Path(__file__).resolve().parent.parent / "temp"
TEMP_ROOT.mkdir(parents=True, exist_ok=True)


def _convert_doc_to_docx(input_path: Path, out_dir: Path) -> Path | None:
    """
    Konverterer .doc → .docx ved hjelp av LibreOffice (soffice).
    Returnerer sti til .docx ved suksess, ellers None.
    """
    try:
        out_dir.mkdir(parents=True, exist_ok=True)
        cmd = ["soffice", "--headless", "--convert-to", "docx", "--outdir", str(out_dir), str(input_path)]
        res = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        if res.returncode != 0:
            log.warning("[.doc→.docx] Konvertering feilet for %s: %s", input_path.name, res.stderr or res.stdout)
            return None
        out_path = out_dir / (input_path.stem + ".docx")
        return out_path if out_path.exists() else None
    except FileNotFoundError:
        log.warning("[.doc→.docx] Fant ikke 'soffice' (LibreOffice). Installer LibreOffice i miljøet for .doc-støtte.")
        return None
    except Exception as e:
        log.warning("[.doc→.docx] Uforutsett feil for %s: %s", input_path.name, e, exc_info=True)
        return None


def _extract_text_from_docx_bytes(data: bytes) -> str:
    """Henter ut tekst fra en .docx-fil gitt som bytes."""
    try:
        doc = Document(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        log.warning("[DOCX-bytes] Kunne ikke lese DOCX-bytes: %s", e)
        return ""


def _pdf_to_text_with_pages(pdf_doc) -> str:
    """Ekstraherer tekst side for side og tagger med [[SIDE n]]."""
    parts = []
    for i, page in enumerate(pdf_doc):
        # 'text' er mest robuste uttrekksmetode for semantikk
        txt = page.get_text("text")
        parts.append(f"[[SIDE {i+1}]]\n{txt}")
    return "\n".join(parts)


def _xlsx_to_text(xlsx_path: Path) -> str:
    """Leser alle celler fra alle ark som tekst (read_only for minnebruk)."""
    try:
        wb = load_workbook(xlsx_path, read_only=True, data_only=True)
    except Exception as e:
        log.error("FEIL ved åpning av XLSX %s: %s", xlsx_path.name, e, exc_info=True)
        return ""
    out = []
    for ws in wb.worksheets:
        out.append(f"[[SIDE 1]]\n")  # XLSX har ikke sider; sett en dummy for pipeline-konsistens
        for row in ws.iter_rows():
            for c in row:
                if c.value is not None:
                    out.append(str(c.value))
    return "\n".join(out)


def _csv_to_text(csv_path: Path) -> str:
    """
    Leser CSV som råtekstlinjer (enkelt og robust).
    NB: Dette er bevisst enkelt; kravsporing jobber på setningsnivå i etterkant.
    """
    try:
        # Prøv UTF-8 først, fall tilbake til latin-1
        try:
            return csv_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return csv_path.read_text(encoding="latin-1", errors="ignore")
    except Exception as e:
        log.error("FEIL ved lesing av CSV %s: %s", csv_path.name, e, exc_info=True)
        return ""


def _html_to_text(s: str) -> str:
    """Veldig enkel HTML→tekst (fjerner tagger og komprimerer whitespace)."""
    if not s:
        return ""
    txt = re.sub(r"<br\s*/?>", "\n", s, flags=re.I)
    txt = re.sub(r"</p\s*>", "\n", txt, flags=re.I)
    txt = re.sub(r"<[^>]+>", " ", txt)
    txt = re.sub(r"\s+", " ", txt).strip()
    return txt


def _process_single_document(
    self_task,
    file_path: Path,
    keywords: list,
    min_score: float,
    ns_standard_selection: str,
    mode: str,
    fokusomraade: str,
    selected_groups: list | None
):
    """
    Leser én enkelt fil (inkl. vedlegg i .msg), trekker ut tekst,
    og kjører kravuthenting på teksten.
    Returnerer en tuple med (funn, feilmeldinger).
    """
    results: list[tuple[list, list[str]]] = []

    def process_text_content(text_content: str, source_name: str, file_type: str):
        """Hjelpefunksjon som kaller selve krav-logikken."""
        file_reqs, file_errs = [], []
        try:
            if text_content:
                cleaned = clean_text(text_content)
                file_reqs = extract_requirements(
                    text=cleaned,
                    selected_function_groups=selected_groups or [],
                    file_name=source_name,
                    min_score=float(min_score),
                    file_type=file_type,
                    ns_standard_selection=ns_standard_selection,
                    mode=mode,
                    fokus_text=fokusomraade or "",
                    use_fokus_prefilter=True,
                    fokus_threshold=0.60,
                    selected_groups=selected_groups or [],
                )
        except Exception as e:
            log.error("FEIL under prosessering av innhold fra %s: %s", source_name, e, exc_info=True)
            file_errs.append(f"Alvorlig feil under prosessering av innhold fra '{source_name}': {e}")
        return file_reqs, file_errs

    filename = file_path.name
    fext = file_path.suffix.lower()

    if fext == ".msg":
        # Les Outlook-epost (tekst + evt. vedlegg)
        try:
            msg = extract_msg.Message(str(file_path))
        except Exception as e:
            log.error("FEIL ved åpning av MSG %s: %s", filename, e, exc_info=True)
            return [], [f"Kritisk feil ved lesing av e-post '{filename}': {e}"]

        # 1) E-post-tekst (body/htmlBody)
        body_txt = msg.body or ""
        try:
            html_txt = getattr(msg, "htmlBody", "") or ""
        except Exception:
            html_txt = ""
        if not body_txt and html_txt:
            body_txt = _html_to_text(html_txt)
        if body_txt:
            reqs, errs = process_text_content(body_txt, f"{filename} (E-post)", file_type="msg")
            results.append((reqs, errs))

        # 2) Vedlegg
        for att in getattr(msg, "attachments", []) or []:
            # extract_msg gir ulike attributter avhengig av versjon; håndter defensivt
            try:
                att_type = getattr(att, "type", "data")
            except Exception:
                att_type = "data"
            if att_type != "data":
                continue

            att_filename = getattr(att, "long_filename", None) or getattr(att, "short_filename", None) or "vedlegg"
            att_ext = Path(att_filename).suffix.lower()
            att_bytes = getattr(att, "data", None)

            if not att_bytes:
                results.append(([], [f"Kunne ikke lese vedlegg '{att_filename}' fra '{filename}' (ingen data)."]))
                continue

            att_content = ""
            try:
                if att_ext == ".pdf":
                    with fitz.open(stream=att_bytes, filetype="pdf") as pdf_doc:
                        att_content = _pdf_to_text_with_pages(pdf_doc)
                    reqs, errs = process_text_content(att_content, f"{filename} -> {att_filename}", file_type="pdf")
                    results.append((reqs, errs))

                elif att_ext == ".docx":
                    att_content = _extract_text_from_docx_bytes(att_bytes)
                    reqs, errs = process_text_content(att_content, f"{filename} -> {att_filename}", file_type="docx")
                    results.append((reqs, errs))

                elif att_ext == ".doc":
                    att_tmp_dir = TEMP_ROOT / f"doc_att_{secrets.token_hex(4)}"
                    att_tmp_dir.mkdir(parents=True, exist_ok=True)
                    try:
                        src_path = att_tmp_dir / Path(att_filename).name
                        src_path.write_bytes(att_bytes)
                        docx_path = _convert_doc_to_docx(src_path, att_tmp_dir)
                        if docx_path and docx_path.exists():
                            att_content = _extract_text_from_docx_bytes(docx_path.read_bytes())
                            reqs, errs = process_text_content(att_content, f"{filename} -> {att_filename}", file_type="doc")
                            results.append((reqs, errs))
                        else:
                            results.append(([], [f"Konvertering av DOC-vedlegg feilet for '{att_filename}'."]))
                    finally:
                        shutil.rmtree(att_tmp_dir, ignore_errors=True)

                elif att_ext in (".txt",):
                    try:
                        att_content = att_bytes.decode("utf-8", errors="ignore")
                    except Exception:
                        att_content = ""
                    if att_content:
                        reqs, errs = process_text_content(att_content, f"{filename} -> {att_filename}", file_type="txt")
                        results.append((reqs, errs))

                elif att_ext in (".csv",):
                    try:
                        # Dekode som utf-8 (fallback latin-1)
                        try:
                            att_content = att_bytes.decode("utf-8")
                        except UnicodeDecodeError:
                            att_content = att_bytes.decode("latin-1", errors="ignore")
                    except Exception:
                        att_content = ""
                    if att_content:
                        reqs, errs = process_text_content(att_content, f"{filename} -> {att_filename}", file_type="csv")
                        results.append((reqs, errs))

                else:
                    # Ukjent eller ikke-støttet vedleggstype – hopp over stille men informer
                    results.append(([], [f"Vedlegg '{att_filename}' (type {att_ext or 'ukjent'}) ble ikke prosessert."]))
            except Exception as e:
                log.warning("Kunne ikke prosessere vedlegg %s: %s", att_filename, e)
                results.append(([], [f"Kunne ikke lese vedlegg '{att_filename}' fra '{filename}'."]))
    else:
        # Håndter enkeltstående filer
        extracted_text = ""
        try:
            if fext == ".pdf":
                with fitz.open(file_path) as doc:
                    extracted_text = _pdf_to_text_with_pages(doc)
                reqs, errs = process_text_content(extracted_text, filename, file_type="pdf")
                results.append((reqs, errs))

            elif fext == ".docx":
                doc = Document(file_path)
                extracted_text = "\n".join(p.text for p in doc.paragraphs)
                reqs, errs = process_text_content(extracted_text, filename, file_type="docx")
                results.append((reqs, errs))

            elif fext == ".doc":
                conv_dir = TEMP_ROOT / f"doc_{secrets.token_hex(4)}"
                conv_dir.mkdir(parents=True, exist_ok=True)
                try:
                    docx_path = _convert_doc_to_docx(file_path, conv_dir)
                    if docx_path and docx_path.exists():
                        extracted_text = _extract_text_from_docx_bytes(docx_path.read_bytes())
                        reqs, errs = process_text_content(extracted_text, filename, file_type="doc")
                        results.append((reqs, errs))
                    else:
                        raise RuntimeError("Konvertering ga ingen .docx-utfil.")
                finally:
                    shutil.rmtree(conv_dir, ignore_errors=True)

            elif fext == ".txt":
                try:
                    extracted_text = file_path.read_text(encoding="utf-8", errors="ignore")
                except Exception as e:
                    log.error("FEIL ved lesing av TXT %s: %s", filename, e, exc_info=True)
                    extracted_text = ""
                reqs, errs = process_text_content(extracted_text, filename, file_type="txt")
                results.append((reqs, errs))

            elif fext == ".xlsx":
                extracted_text = _xlsx_to_text(file_path)
                reqs, errs = process_text_content(extracted_text, filename, file_type="xlsx")
                results.append((reqs, errs))

            elif fext == ".csv":
                extracted_text = _csv_to_text(file_path)
                reqs, errs = process_text_content(extracted_text, filename, file_type="csv")
                results.append((reqs, errs))

            else:
                results.append(([], [f"Filtype {fext or 'ukjent'} støttes ikke for '{filename}'."]))

        except Exception as e:
            log.error("FEIL ved lesing av fil %s: %s", filename, e, exc_info=True)
            results.append(([], [f"Kritisk feil ved lesing av fil '{filename}': {e}"]))

    final_reqs = [req for res_tuple in results for req in (res_tuple[0] or [])]
    final_errs = [err for res_tuple in results for err in (res_tuple[1] or [])]

    return final_reqs, final_errs
