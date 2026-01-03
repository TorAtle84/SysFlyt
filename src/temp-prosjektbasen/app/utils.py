from flask import Blueprint, request, jsonify, Response, current_app
from flask_login import login_required
import hashlib
import os, re, logging, time, json
from pathlib import Path
from hashlib import md5
import colorsys
import pdfplumber
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from collections import defaultdict
import fitz  # PyMuPDF
from openpyxl import Workbook, load_workbook
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from io import BytesIO
from PIL import Image, ImageDraw
import pytesseract
from rapidfuzz.fuzz import partial_ratio as fuzz_partial_ratio
from ics import Calendar, Event, Organizer, Attendee # VIKTIG: Sørg for disse er importert
import uuid 
from app.models.db import db # Antar 'db' er korrekt importert fra 'app' instansen
from app.models.task import Task # Trengs for TaskRevision.task relasjon
import tempfile 
import shutil   
import os  
from datetime import datetime, timezone, timedelta

# Note: oppgave_api Blueprint definition here is likely a copy-paste error from oppgaver.py.
# Blueprints should be defined only once in their respective files.
# If you copied this from oppgaver.py, please remove it from utils.py.
# oppgave_api = Blueprint("oppgave_api", __name__) 

# === Fra før: farge for tekniker ===
def generer_farge_for_tekniker(navn):
    h = hashlib.md5(navn.encode()).hexdigest()
    r = int(h[:2], 16)
    g = int(h[2:4], 16)
    b = int(h[4:6], 16)
    return f"rgba({r},{g},{b},0.95)"
    
LOKASJONER = [
    "Bergen", "Oslo", "Trondheim", "Region Vest (øvrig)",
    "Region Nord (øvrig)", "Region Sør (øvrig)", "Region Øst (øvrig)"
]

FAG = ["ventilasjon", "byggautomasjon", "rør", "elektro", "kulde"]

def ensure_ics_uid(task):
    if not getattr(task, "ics_uid", None):
        task.ics_uid = str(uuid.uuid4())
        current_app.logger.debug(f"[ensure_ics_uid] Sett ny ics_uid: {task.ics_uid} for task {task.id}")

def build_ics_string(ev: Event, method: str, organizer_email: str = None, attendee_email: str = None) -> str:
    cal = Calendar()
    cal.prodid = "-//GK//TeknikerBooking//NO"
    cal.created = datetime.now(timezone.utc)
    if organizer_email: ev.organizer = Organizer(f"MAILTO:{organizer_email}", common_name=organizer_email.split('@')[0].capitalize())
    if attendee_email:
        if not isinstance(ev.attendees, list): ev.attendees = []
        ev.attendees.append(Attendee(f"MAILTO:{attendee_email}", common_name=attendee_email.split('@')[0].capitalize()))
    cal.events.add(ev)
    ics_text_raw = cal.serialize()
    lines = ics_text_raw.splitlines()
    fixed_lines = []
    inserted_method = False
    for line in lines:
        fixed_lines.append(line)
        if line.strip() == "BEGIN:VCALENDAR" and not inserted_method:
            fixed_lines.append(f"METHOD:{method}")
            inserted_method = True
    final_ics_string = "\r\n".join(fixed_lines) + "\r\n"
    return final_ics_string

def send_ics_cancel(ics_uid: str, ics_sequence: int, old_title: str, old_start_date: datetime.date, old_start_time: datetime.time, old_end_date: datetime.date, old_end_time: datetime.time, organizer_email: str = None, attendee_email: str = None) -> str:
    ev = Event()
    ev.uid = ics_uid
    ev.sequence = ics_sequence
    ev.status = "CANCELLED"
    ev.method = "CANCEL"
    ev.name = old_title
    ev.created = datetime.now(timezone.utc)
    if old_start_time is None: ev.begin = old_start_date.isoformat()
    else: ev.begin = datetime.combine(old_start_date, old_start_time, tzinfo=timezone.utc)
    if old_end_time is None: ev.end = (old_end_date + timedelta(days=1)).isoformat()
    else: ev.end = datetime.combine(old_end_date, old_end_time, tzinfo=timezone.utc)
    ics_string_output = build_ics_string(ev, "CANCEL", organizer_email, attendee_email)
    current_app.logger.info("GENERATED CANCEL ICS (UID=%s, SEQ=%d):\n%s", ics_uid, ics_sequence, ics_string_output)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".ics")
    tmp.write(ics_string_output.encode("utf-8"))
    tmp.close()
    return tmp.name
    
def generate_dynamic_pattern(example):
    """Lager regex-mønster fra eksempel."""
    pattern = ""
    for char in example:
        if char.isalpha():
            pattern += "[A-Za-z]"
        elif char.isdigit():
            pattern += r"\d"
        elif char in [".", ","]:
            pattern += "[.,]"
        elif char in ["-", "\u2013", "_"]:
            pattern += "[-\u2013_]?"
        else:
            pattern += re.escape(char)
    return pattern

def extract_keywords(file_path: str, pattern: str, use_ocr: bool = False) -> set:
    """Enkel keyword-extract for dokumentsammenligning."""
    import fitz
    from docx import Document
    import pdfplumber
    from PIL import Image
    import pytesseract

    ext = Path(file_path).suffix.lower()
    text = ""

    if ext == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            text = "\n".join([page.extract_text() or "" for page in pdf.pages])
    elif ext == ".docx":
        doc = Document(file_path)
        text = "\n".join([p.text for p in doc.paragraphs])
    elif ext == ".txt":
        text = Path(file_path).read_text(encoding="utf-8")
    else:
        text = ""

    matches = re.findall(pattern, text)
    return set(matches)

def generate_report(hoveddokument, sammenligningsfiler, keyword_pattern, use_hoveddokument_keywords_only, use_ocr=False):
    """Lager sammenligningsrapport for nøkkelord mellom dokumenter."""
    pattern = generate_dynamic_pattern(keyword_pattern)
    hoveddokument_keywords = extract_keywords(hoveddokument, pattern, use_ocr)

    sammenlignings_keywords = {}
    for fil in sammenligningsfiler:
        key = Path(fil).stem
        sammenlignings_keywords[key] = extract_keywords(fil, pattern, use_ocr)

    all_keywords = (
        sorted(hoveddokument_keywords) if use_hoveddokument_keywords_only
        else sorted(set(hoveddokument_keywords).union(*sammenlignings_keywords.values()))
    )

    wb = Workbook()
    ws = wb.active
    ws.title = "Sammenligning"

    headers = ["Nøkkelord", "Hoveddokument"] + list(sammenlignings_keywords.keys())
    ws.append(headers)

    header_fill = PatternFill(start_color="FFD3D3D3", end_color="FFD3D3D3", fill_type="solid")
    bold_font = Font(bold=True)
    thin_border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin")
    )

    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = bold_font
        cell.border = thin_border

    true_fill = PatternFill(start_color="FFA4FFA4", end_color="FFA4FFA4", fill_type="solid")
    false_fill = PatternFill(start_color="FFFF7E8E", end_color="FFFF7E8E", fill_type="solid")

    for keyword in all_keywords:
        row = [keyword]
        row.append("Ja" if keyword in hoveddokument_keywords else "Nei")
        for doc_keywords in sammenlignings_keywords.values():
            row.append("Ja" if keyword in doc_keywords else "Nei")
        ws.append(row)

        for col_idx, value in enumerate(row[1:], start=2):
            cell = ws.cell(row=ws.max_row, column=col_idx)
            cell.fill = true_fill if value == "Ja" else false_fill
            cell.border = thin_border

    ws.auto_filter.ref = ws.dimensions

    for col in ws.columns:
        max_len = max((len(str(cell.value)) for cell in col), default=10)
        ws.column_dimensions[get_column_letter(col[0].column)].width = min(max_len + 5, 50)

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    return output
    
def extract_text_from_file(file_storage):
    import io, docx, fitz, openpyxl
    file_storage.seek(0)
    file_bytes = file_storage.read()
    filename = file_storage.filename.lower()
    content = ""

    if filename.endswith(".pdf"):
        pdf = fitz.open(stream=file_bytes, filetype="pdf")
        for page in pdf:
            content += page.get_text()
    elif filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            content += para.text + "\n"
    elif filename.endswith(".xlsx"):
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                content += " ".join([str(cell) for cell in row if cell is not None]) + "\n"
    elif filename.endswith(".txt"):
        content = file_bytes.decode("utf-8", errors="ignore")
    return content

def extract_system_numbers(content, pattern, prefix, system_start):
    import re
    regex = re.compile(pattern)
    matches = regex.findall(content)
    return [
        m for m in matches
        if m.startswith(prefix) and m[len(prefix):].startswith(system_start or "")
    ]